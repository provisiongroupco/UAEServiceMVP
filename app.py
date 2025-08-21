import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import os
from PIL import Image
from streamlit_drawable_canvas import st_canvas
import numpy as np
import json
import base64
import binascii
import urllib.parse
from utils import (style_heading, create_info_table, set_cell_margins, format_table_style_enhanced)
from equipment_config import EQUIPMENT_TYPES, MARVEL_CHECKLIST

# Page configuration
st.set_page_config(
    page_title="Service Reports System",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
    <style>
    .main-header {
        font-size: 2.2rem;
        color: #1f4788;
        font-weight: bold;
        text-align: center;
        margin-bottom: 1.2rem;
    }
    .section-header {
        font-size: 1.4rem;
        color: #2c5aa0;
        font-weight: bold;
        margin-top: 0.8rem;
        margin-bottom: 0.6rem;
    }
    .stButton > button {
        background-color: #1f4788;
        color: white;
        font-weight: bold;
        padding: 0.5rem 2rem;
        border-radius: 5px;
        border: none;
        width: 100%;
    }
    .stButton > button:hover {
        background-color: #2c5aa0;
    }
    /* Reduce spacing between elements */
    .stTextInput > div > div > input {
        margin-bottom: 0;
    }
    .stSelectbox > div > div {
        margin-bottom: 0;
    }
    .stTextArea > div > div > textarea {
        margin-bottom: 0;
    }
    div[data-testid="stVerticalBlock"] > div {
        gap: 0.8rem;
    }
    .element-container {
        margin-bottom: 0.5rem;
    }
    h3 {
        margin-top: 1rem !important;
        margin-bottom: 0.5rem !important;
    }
    h4 {
        margin-top: 0.8rem !important;
        margin-bottom: 0.4rem !important;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
if 'report_data' not in st.session_state:
    st.session_state.report_data = {}
if 'report_generated' not in st.session_state:
    st.session_state.report_generated = False
if 'technician_signature' not in st.session_state:
    st.session_state.technician_signature = None
if 'kitchen_list' not in st.session_state:
    st.session_state.kitchen_list = []
if 'form_data' not in st.session_state:
    st.session_state.form_data = {}

# K-Factor lookup tables based on PDF documentation
def get_extract_k_factor(num_filters, hood_type, with_uv=False):
    """Get K-Factor for extract air based on number of KSA filters and hood type"""
    # K-Factor tables in m³/h
    if with_uv:
        # UVF/UVI hoods with UV
        uv_k_factors = {
            1: 53.82,
            2: 107.64,
            3: 161.46,
            4: 215.28,
            5: 269.1,
            6: 322.92
        }
        return uv_k_factors.get(num_filters, 0)
    else:
        # KVF/KVI standard hoods without UV
        standard_k_factors = {
            1: 67.21,
            2: 134.46,
            3: 201.67,
            4: 268.88,
            5: 336.09,
            6: 403.30
        }
        return standard_k_factors.get(num_filters, 0)

def get_supply_k_factor(hood_length):
    """Get K-Factor for supply air based on hood length (H-555 model)"""
    # K-Factor table for supply air (hood length in meters)
    # Linear interpolation for values between table entries
    supply_k_table = [
        (1.0, 121.7),
        (1.1, 133.9),
        (1.2, 146.1),
        (1.3, 158.2),
        (1.4, 170.4),
        (1.5, 182.6),
        (1.6, 194.8),
        (1.7, 207.0),
        (1.8, 219.1),
        (1.9, 231.3),
        (2.0, 243.3),
        (2.1, 255.5),
        (2.2, 267.7),
        (2.3, 279.9),
        (2.4, 292.0),
        (2.5, 304.2),
        (2.6, 316.4),
        (2.7, 328.6),
        (2.8, 340.8),
        (2.9, 352.9),
        (3.0, 365.1),
        (3.1, 377.3),
        (3.2, 389.5),
        (3.3, 401.7),
        (3.4, 413.8),
        (3.5, 426.0),
        (3.6, 438.2),
        (3.7, 450.4),
        (3.8, 462.6),
        (3.9, 474.7),
        (4.0, 486.9)
    ]
    
    # Find the appropriate K-factor
    if hood_length <= 1.0:
        return 121.7
    elif hood_length >= 4.0:
        return 486.9
    else:
        # Linear interpolation between values
        for i in range(len(supply_k_table) - 1):
            if supply_k_table[i][0] <= hood_length <= supply_k_table[i+1][0]:
                x1, y1 = supply_k_table[i]
                x2, y2 = supply_k_table[i+1]
                # Linear interpolation
                k_factor = y1 + (y2 - y1) * (hood_length - x1) / (x2 - x1)
                return round(k_factor, 1)
    return 0

def render_checklist_item(equipment, item, equip_key_prefix, prefix=""):
    """Recursively render a checklist item with all its conditional logic"""
    item_key = prefix + item['id']
    
    # Ensure inspection_data exists
    if 'inspection_data' not in equipment:
        equipment['inspection_data'] = {}
    
    if item_key not in equipment['inspection_data']:
        equipment['inspection_data'][item_key] = {}
    
    item_data = equipment['inspection_data'][item_key]
    
    # Render the question
    question = item['question']
    question_type = item['type']
    
    if question_type == 'yes_no':
        options = ['', 'Yes', 'No']
        widget_key = f"q_{item_key}_{equip_key_prefix}"
        
        # Initialize session state if not exists
        if widget_key not in st.session_state:
            st.session_state[widget_key] = item_data.get('answer', '')
            
        answer = st.selectbox(
            question,
            options=options,
            key=widget_key
        )
        # Update equipment object directly
        equipment['inspection_data'][item_key]['answer'] = st.session_state[widget_key]
        
    elif question_type == 'yes_no_na':
        options = ['', 'Yes', 'No', 'N/A']
        widget_key = f"q_{item_key}_{equip_key_prefix}"
        
        # Initialize session state if not exists
        if widget_key not in st.session_state:
            st.session_state[widget_key] = item_data.get('answer', '')
            
        answer = st.selectbox(
            question,
            options=options,
            key=widget_key
        )
        # Update equipment object directly
        equipment['inspection_data'][item_key]['answer'] = st.session_state[widget_key]
        
    elif question_type == 'text':
        widget_key = f"q_{item_key}_{equip_key_prefix}"
        
        # Initialize session state if not exists
        if widget_key not in st.session_state:
            st.session_state[widget_key] = item_data.get('answer', '')
        
        # Use text_area for dimensions questions to provide more space
        if 'dimensions' in item_key.lower() or 'sizes' in question.lower():
            answer = st.text_area(
                question,
                key=widget_key,
                height=100  # Provides a larger text box
            )
        else:
            answer = st.text_input(
                question,
                key=widget_key
            )
        # Update equipment object directly
        equipment['inspection_data'][item_key]['answer'] = st.session_state[widget_key]
        
        # Handle photo requirement for text fields
        if answer and item.get('photo'):
            photo_key = f"photo_{item_key}"
            uploaded_files = st.file_uploader(
                f"📷 Upload photo(s) for: {question}",
                type=['png', 'jpg', 'jpeg'],
                key=f"photo_{item_key}_{equip_key_prefix}",
                accept_multiple_files=True
            )
            if uploaded_files:
                if 'photos' not in equipment:
                    equipment['photos'] = {}
                # Store multiple photos with numbered keys
                for i, uploaded_file in enumerate(uploaded_files):
                    if len(uploaded_files) == 1:
                        equipment['photos'][photo_key] = uploaded_file
                    else:
                        equipment['photos'][f"{photo_key}_{i+1}"] = uploaded_file
                st.success(f"✅ {len(uploaded_files)} photo(s) uploaded")
        
    elif question_type == 'number':
        widget_key = f"q_{item_key}_{equip_key_prefix}"
        
        # Initialize session state if not exists
        if widget_key not in st.session_state:
            st.session_state[widget_key] = item_data.get('answer', 0)
            
        answer = st.number_input(
            question,
            min_value=0,
            step=1,
            key=widget_key
        )
        # Update equipment object directly
        equipment['inspection_data'][item_key]['answer'] = st.session_state[widget_key]
        
        # Handle alarm generation if this is an alarm count question
        if item.get('generates_alarms') and answer > 0:
            st.markdown("##### Alarm Details")
            
            # Initialize alarm data structure if not exists
            if 'alarm_details' not in equipment:
                equipment['alarm_details'] = {}
            
            # Generate alarm detail fields for each alarm
            for i in range(int(answer)):
                alarm_idx = i + 1
                alarm_key = f"alarm_{alarm_idx}"
                
                # Initialize alarm data if not exists
                if alarm_key not in equipment['alarm_details']:
                    equipment['alarm_details'][alarm_key] = {}
                
                st.markdown(f"**Alarm {alarm_idx}:**")
                
                # Description field
                desc_key = f"alarm_desc_{alarm_idx}_{equip_key_prefix}"
                if desc_key not in st.session_state:
                    st.session_state[desc_key] = equipment['alarm_details'][alarm_key].get('description', '')
                
                description = st.text_area(
                    f"Description for Alarm {alarm_idx}",
                    key=desc_key,
                    height=80,
                    placeholder="Describe the alarm (e.g., UV lamp failure, Communication error, etc.)"
                )
                equipment['alarm_details'][alarm_key]['description'] = st.session_state[desc_key]
                
                # Photo upload for this alarm
                alarm_photo_key = f"photo_alarm_{alarm_idx}"
                uploaded_files = st.file_uploader(
                    f"📷 Upload photo(s) for Alarm {alarm_idx}",
                    type=['png', 'jpg', 'jpeg'],
                    key=f"photo_alarm_{alarm_idx}_{equip_key_prefix}",
                    accept_multiple_files=True
                )
                
                if uploaded_files:
                    if 'photos' not in equipment:
                        equipment['photos'] = {}
                    # Store multiple photos with numbered keys
                    for j, uploaded_file in enumerate(uploaded_files):
                        if len(uploaded_files) == 1:
                            equipment['photos'][alarm_photo_key] = uploaded_file
                        else:
                            equipment['photos'][f"{alarm_photo_key}_{j+1}"] = uploaded_file
                    st.success(f"✅ {len(uploaded_files)} photo(s) uploaded for Alarm {alarm_idx}")
                
                # Add separator between alarms
                if i < int(answer) - 1:
                    st.markdown("---")
    else:
        answer = None
    
    # Get answer from updated data for conditional logic
    answer = equipment['inspection_data'][item_key].get('answer', '')
    
    # Handle conditional logic based on answer
    if answer and 'conditions' in item:
        # Convert answer to lowercase for matching
        answer_key = answer.lower()
        condition = item['conditions'].get(answer_key)
        if condition:
            # Show a visual indicator that this answer triggered conditions
            st.markdown(f"<small style='color: #1f4788; margin: 0;'>→ Based on your answer '{answer}':</small>", unsafe_allow_html=True)
            
            # Handle photo requirement
            if condition.get('photo'):
                photo_key = f"photo_{item_key}"
                uploaded_files = st.file_uploader(
                    f"📷 Upload photo(s) for: {question}",
                    type=['png', 'jpg', 'jpeg'],
                    key=f"photo_{item_key}_{equip_key_prefix}",
                    accept_multiple_files=True
                )
                if uploaded_files:
                    if 'photos' not in equipment:
                        equipment['photos'] = {}
                    # Store multiple photos with numbered keys
                    for i, uploaded_file in enumerate(uploaded_files):
                        if len(uploaded_files) == 1:
                            equipment['photos'][photo_key] = uploaded_file
                        else:
                            equipment['photos'][f"{photo_key}_{i+1}"] = uploaded_file
                    st.success(f"✅ {len(uploaded_files)} photo(s) uploaded")
            
            # Handle comment requirement
            if condition.get('comment'):
                comment_key = f"comment_{item_key}_{equip_key_prefix}"
                
                # Initialize session state if not exists
                if comment_key not in st.session_state:
                    st.session_state[comment_key] = item_data.get('comment', '')
                    
                comment = st.text_area(
                    f"Please provide details for: {question}",
                    key=comment_key,
                    height=100
                )
                equipment['inspection_data'][item_key]['comment'] = st.session_state[comment_key]
            
            # Handle action instruction
            if condition.get('action'):
                st.warning(condition['action'])
            
            # Handle follow-up questions
            if condition.get('follow_up'):
                # Add indentation for follow-up questions
                with st.container():
                    # Create a container with padding for visual hierarchy
                    _, col2 = st.columns([0.05, 0.95])
                    with col2:
                        for follow_up_item in condition['follow_up']:
                            render_checklist_item(equipment, follow_up_item, equip_key_prefix, prefix)


def find_question_text(equipment_type, item_key):
    """Find the actual question text for a given item key"""
    def search_checklist(checklist_items, target_id):
        for item in checklist_items:
            if item['id'] == target_id:
                return item['question']
            # Check in follow-up questions
            if 'conditions' in item:
                for condition in item['conditions'].values():
                    if 'follow_up' in condition:
                        result = search_checklist(condition['follow_up'], target_id)
                        if result:
                            return result
        return None
    
    # Check if this is a Marvel question
    if item_key.startswith('marvel_'):
        # Remove marvel_ prefix and search in MARVEL checklist
        marvel_key = item_key[7:]  # Remove 'marvel_' prefix
        question = search_checklist(MARVEL_CHECKLIST, marvel_key)
        if question:
            return question
    
    # Get the last part of the key which is the actual question ID
    parts = item_key.split('_')
    for i in range(len(parts)):
        # Try different combinations starting from the end
        for j in range(len(parts), i, -1):
            test_id = '_'.join(parts[i:j])
            question = search_checklist(EQUIPMENT_TYPES[equipment_type]['checklist'], test_id)
            if question:
                return question
    
    # Fallback to formatted key
    return item_key.replace('_', ' ').title()


def get_kitchen_summary():
    """Get a summary of all kitchen and equipment inspections"""
    summary = []
    
    for kitchen in st.session_state.kitchen_list:
        kitchen_summary = {
            'name': kitchen.get('name', 'Unknown Kitchen'),
            'equipment': []
        }
        
        for equipment in kitchen.get('equipment_list', []):
            if equipment.get('type'):  # Only include equipment with a selected type
                equip_summary = {
                    'type': equipment['type'],
                    'type_name': EQUIPMENT_TYPES[equipment['type']]['name'],
                    'with_marvel': equipment.get('with_marvel', False),
                    'location': equipment.get('location', ''),
                    'yes_responses': [],
                    'no_responses': [],
                    'na_responses': [],
                    'photos_count': len(equipment.get('photos', {})),
                    'inspection_data': equipment.get('inspection_data', {}),
                    'photos': equipment.get('photos', {}),
                    'yes_photos': {},
                    'no_photos': {},
                    'na_photos': {},
                    'alarm_details': equipment.get('alarm_details', {})
                }
                
                # Define items to exclude from No responses (these are not issues)
                exclude_from_no = ['final_remarks', 'lights_ballast']
            
                # Separate Yes and No responses
                for key, data in equipment.get('inspection_data', {}).items():
                    if isinstance(data, dict):
                        answer = data.get('answer', '')
                        
                        # Skip Marvel questions if Marvel is not enabled
                        if key.startswith('marvel_') and not equipment.get('with_marvel', False):
                            continue
                        
                        # Skip questions that don't belong to current equipment type
                        if not key.startswith('marvel_'):
                            # Check if this question belongs to the current equipment type
                            def check_question_in_checklist(checklist_items, target_key):
                                for item in checklist_items:
                                    if item['id'] in target_key:
                                        return True
                                    # Check in follow-up questions
                                    if 'conditions' in item:
                                        for condition in item['conditions'].values():
                                            if 'follow_up' in condition:
                                                if check_question_in_checklist(condition['follow_up'], target_key):
                                                    return True
                                return False
                            
                            question_found = False
                            if equipment.get('type'):
                                checklist = EQUIPMENT_TYPES.get(equipment['type'], {}).get('checklist', [])
                                question_found = check_question_in_checklist(checklist, key)
                            
                            if not question_found:
                                continue  # Skip this question as it doesn't belong to current equipment
                        
                        # Extract the base key - get the actual question ID
                        # Keys might be like "lights_ballast" or "lights_ballast_ballast_issue"
                        # We want to check if any part of the key is in the exclude list
                        key_parts = key.split('_')
                        should_exclude = False
                        for i in range(len(key_parts)):
                            check_key = '_'.join(key_parts[i:])
                            if check_key in exclude_from_no:
                                should_exclude = True
                                break
                        
                        # Get the actual question text
                        question_text = find_question_text(equipment['type'], key)
                        
                        if answer == 'Yes':
                            equip_summary['yes_responses'].append({
                                'item': key,
                                'question': question_text,
                                'answer': answer,
                                'comment': data.get('comment', '')
                            })
                            # Collect photos for this Yes response
                            photo_key = f"photo_{key}"
                            for pk, pv in equipment.get('photos', {}).items():
                                if pk.startswith(photo_key):
                                    equip_summary['yes_photos'][pk] = pv
                        elif answer == 'No':
                            # Only add to no_responses if it's not in the exclude list
                            if not should_exclude:
                                equip_summary['no_responses'].append({
                                    'item': key,
                                    'question': question_text,
                                    'answer': answer,
                                    'comment': data.get('comment', '')
                                })
                                # Collect photos for this No response
                                photo_key = f"photo_{key}"
                                for pk, pv in equipment.get('photos', {}).items():
                                    if pk.startswith(photo_key):
                                        equip_summary['no_photos'][pk] = pv
                        elif answer == 'N/A':
                            equip_summary['na_responses'].append({
                                'item': key,
                                'question': question_text,
                                'answer': answer,
                                'comment': data.get('comment', '')
                            })
                            # Collect photos for this N/A response (though N/A usually doesn't have photos)
                            photo_key = f"photo_{key}"
                            for pk, pv in equipment.get('photos', {}).items():
                                if pk.startswith(photo_key):
                                    equip_summary['na_photos'][pk] = pv
                        elif answer and answer not in ['', '0']:  # Text or number responses
                            # Add to yes_responses for display purposes
                            equip_summary['yes_responses'].append({
                                'item': key,
                                'question': question_text,
                                'answer': answer,
                                'comment': data.get('comment', '')
                            })
                            # Collect photos for text/number responses
                            photo_key = f"photo_{key}"
                            for pk, pv in equipment.get('photos', {}).items():
                                if pk.startswith(photo_key):
                                    equip_summary['yes_photos'][pk] = pv
                    
                # For backward compatibility, keep issues_found as no_responses
                equip_summary['issues_found'] = equip_summary['no_responses']
                
                kitchen_summary['equipment'].append(equip_summary)
        
        if kitchen_summary['equipment']:  # Only add kitchen if it has equipment
            summary.append(kitchen_summary)
    
    return summary


def collect_form_data():
    """Collect all form data from session state for sharing"""
    try:
        form_data = {
            'basic_info': {
                'customer_name': st.session_state.get('customer_name', ''),
                'project_name': st.session_state.get('project_name', ''),
                'contact_person': st.session_state.get('contact_person', ''),
                'outlet_location': st.session_state.get('outlet_location', ''),
                'contact_number': st.session_state.get('contact_number', ''),
                'visit_type': st.session_state.get('visit_type', ''),
                'visit_class': st.session_state.get('visit_class', ''),
                'report_date': st.session_state.get('report_date', datetime.now()).isoformat() if st.session_state.get('report_date') else None,
                'work_performed': st.session_state.get('work_performed', ''),
                'spare_parts': st.session_state.get('spare_parts', []),
                'recommendations': st.session_state.get('recommendations', ''),
                'technician_name': st.session_state.get('technician_name', ''),
                'service_date': st.session_state.get('service_date', datetime.now()).isoformat() if st.session_state.get('service_date') else None,
                'report_type': st.session_state.get('report_type', 'Technical Report'),
                'work_performed_list': st.session_state.get('work_performed_list', [])
            },
            'kitchen_data': {
                'num_kitchens': st.session_state.get('num_kitchens', 1),
                'kitchen_list': []
            }
        }
        
        # Collect kitchen and equipment data (excluding photos and signatures for size reasons)
        for kitchen in st.session_state.get('kitchen_list', []):
            kitchen_data = {
                'name': kitchen.get('name', ''),
                'equipment_list': []
            }
            
            for equipment in kitchen.get('equipment_list', []):
                equipment_data = {
                    'type': equipment.get('type', ''),
                    'with_marvel': equipment.get('with_marvel', False),
                    'location': equipment.get('location', ''),
                    'inspection_data': {},
                    'alarm_details': equipment.get('alarm_details', {})
                }
                
                # Include inspection answers but exclude photos
                for key, data in equipment.get('inspection_data', {}).items():
                    if isinstance(data, dict):
                        equipment_data['inspection_data'][key] = {
                            'answer': data.get('answer', ''),
                            'comment': data.get('comment', '')
                        }
                
                kitchen_data['equipment_list'].append(equipment_data)
            
            form_data['kitchen_data']['kitchen_list'].append(kitchen_data)
    
        return form_data
    except Exception as e:
        st.error(f"Error collecting form data: {str(e)}")
        return None


def encode_form_data_to_url(form_data):
    """Encode form data to a URL-safe string"""
    try:
        # Convert to JSON and compress
        json_str = json.dumps(form_data, separators=(',', ':'))
        
        # Base64 encode for URL safety
        encoded_bytes = base64.urlsafe_b64encode(json_str.encode('utf-8'))
        encoded_str = encoded_bytes.decode('utf-8')
        
        # URL encode for additional safety
        url_encoded = urllib.parse.quote(encoded_str)
        
        return url_encoded
    except Exception as e:
        st.error(f"Error encoding form data: {str(e)}")
        return None


def decode_form_data_from_url(encoded_data):
    """Decode form data from URL parameter"""
    try:
        # Basic validation
        if not encoded_data or len(encoded_data) < 10:
            st.error("Invalid share link: Data parameter is too short")
            return None
        
        if len(encoded_data) > 10000:  # Reasonable limit for URL length
            st.error("Invalid share link: Data parameter is too long")
            return None
        
        # URL decode
        url_decoded = urllib.parse.unquote(encoded_data)
        
        # Base64 decode
        decoded_bytes = base64.urlsafe_b64decode(url_decoded.encode('utf-8'))
        json_str = decoded_bytes.decode('utf-8')
        
        # Parse JSON
        form_data = json.loads(json_str)
        
        # Validate structure
        if not isinstance(form_data, dict):
            st.error("Invalid share link: Data format is incorrect")
            return None
            
        if 'basic_info' not in form_data or 'kitchen_data' not in form_data:
            st.error("Invalid share link: Missing required data sections")
            return None
        
        return form_data
    except base64.binascii.Error:
        st.error("Invalid share link: Data encoding is corrupted")
        return None
    except json.JSONDecodeError:
        st.error("Invalid share link: Data format is corrupted")
        return None
    except Exception as e:
        st.error(f"Error decoding share link: {str(e)}")
        return None


def restore_form_data(form_data):
    """Restore form data to session state"""
    try:
        # Debug: Show what we're restoring
        st.sidebar.write("🔄 Restoring data to session state...")
        
        # Restore basic info
        basic_info = form_data.get('basic_info', {})
        restored_count = 0
        for key, value in basic_info.items():
            if value:  # Only restore non-empty values
                if key in ['report_date', 'service_date'] and value:
                    # Parse ISO date strings back to datetime objects
                    st.session_state[key] = datetime.fromisoformat(value).date()
                else:
                    st.session_state[key] = value
                restored_count += 1
        
        st.sidebar.write(f"📝 Restored {restored_count} basic info fields")
        
        # Restore spare parts with proper counter initialization
        if 'spare_parts' in basic_info and basic_info['spare_parts']:
            st.session_state['spare_parts'] = basic_info['spare_parts']
            # Initialize counter and ensure all parts have unique IDs
            import time
            for i, part in enumerate(st.session_state['spare_parts']):
                if 'id' not in part or not part['id']:
                    # Add unique ID if missing
                    part['id'] = f"restored_{i}_{int(time.time() * 1000)}"
            st.session_state['spare_parts_counter'] = len(st.session_state['spare_parts'])
        
        # Restore report type
        if 'report_type' in basic_info:
            st.session_state['report_type'] = basic_info['report_type']
        
        # Restore work performed list for General Service Report
        if 'work_performed_list' in basic_info and basic_info['work_performed_list']:
            st.session_state['work_performed_list'] = basic_info['work_performed_list']
            # Initialize counter and ensure all work items have unique IDs
            for i, work_item in enumerate(st.session_state['work_performed_list']):
                if 'id' not in work_item or not work_item['id']:
                    # Add unique ID if missing
                    work_item['id'] = f"restored_work_{i}_{int(time.time() * 1000)}"
            st.session_state['work_performed_counter'] = len(st.session_state['work_performed_list'])
        
        # Restore kitchen data
        kitchen_data = form_data.get('kitchen_data', {})
        if kitchen_data.get('num_kitchens'):
            st.session_state['num_kitchens'] = kitchen_data['num_kitchens']
        
        # Restore kitchen list
        if kitchen_data.get('kitchen_list'):
            st.session_state.kitchen_list = []
            for kitchen_info in kitchen_data['kitchen_list']:
                kitchen = {
                    'id': f"kitchen_{len(st.session_state.kitchen_list)}_{datetime.now().strftime('%Y%m%d%H%M%S')}",
                    'name': kitchen_info.get('name', ''),
                    'equipment_list': []
                }
                
                for equipment_info in kitchen_info.get('equipment_list', []):
                    equipment = {
                        'id': f"equipment_{len(kitchen['equipment_list'])}_{datetime.now().strftime('%Y%m%d%H%M%S')}",
                        'type': equipment_info.get('type', ''),
                        'with_marvel': equipment_info.get('with_marvel', False),
                        'location': equipment_info.get('location', ''),
                        'inspection_data': equipment_info.get('inspection_data', {}),
                        'alarm_details': equipment_info.get('alarm_details', {}),
                        'photos': {}  # Photos are not shared via links
                    }
                    kitchen['equipment_list'].append(equipment)
                
                st.session_state.kitchen_list.append(kitchen)
        
        return True
    except Exception as e:
        st.error(f"Error restoring form data: {str(e)}")
        return False


def generate_shareable_link():
    """Generate a shareable link with current form data"""
    form_data = collect_form_data()
    if not form_data:
        st.error("Failed to collect form data for sharing")
        return None
        
    encoded_data = encode_form_data_to_url(form_data)
    
    if encoded_data:
        # Use the actual Streamlit app URL
        base_url = "https://ksaservicemvp.streamlit.app"
        share_url = f"{base_url}?data={encoded_data}"
        return share_url
    return None


def create_technical_report(data):
    """Generate a Professional Technical Report Word document"""
    
    # Helper function to set cell background
    def set_cell_background(cell, color):
        """Set background color for a table cell"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)
    
    # Helper function to set table borders
    def set_table_borders(table):
        """Apply borders to table"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Remove existing borders
        for child in tblPr:
            if child.tag.endswith('tblBorders'):
                tblPr.remove(child)
        
        # Create new borders
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    # Use template if available, otherwise create new document
    template_path = "Templates/Report Letter Head.docx"
    
    try:
        if os.path.exists(template_path):
            doc = Document(template_path)
            # Template already has margins and header/footer set up
        else:
            # Fallback to creating a new document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
                section.header_distance = Inches(0.5)
                section.footer_distance = Inches(0.5)
    except Exception as e:
        st.warning(f"Could not load template: {str(e)}. Using blank document.")
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)
    
    # Add some initial spacing
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add title manually to avoid underline
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('TECHNICAL REPORT')
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(31, 71, 136)  # Professional Blue
    title_run.font.bold = True
    
    # Add report reference and date
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref_para.add_run(f"Report Date: {data.get('date', datetime.now().strftime('%B %d, %Y'))}")
    ref_run.font.size = Pt(11)
    ref_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Add minimal spacing
    doc.add_paragraph()
    
    # GENERAL INFORMATION SECTION
    general_heading = doc.add_heading('1. GENERAL INFORMATION', level=1)
    style_heading(general_heading, level=1)
    
    # Create professional info table
    general_info = [
        ("Customer Name", data.get('customer_name', '')),
        ("Project Name", data.get('project_name', '')),
        ("Contact Person", data.get('contact_person', '')),
        ("Location", data.get('outlet_location', '')),
        ("Contact Number", data.get('contact_number', '')),
        ("Visit Type", data.get('visit_type', '')),
        ("Visit Classification", data.get('visit_class', ''))
    ]
    
    create_info_table(doc, general_info)
    doc.add_paragraph()  # Add spacing
    
    # EQUIPMENT INSPECTION SECTION
    equipment_heading = doc.add_heading('2. EQUIPMENT INSPECTION DETAILS', level=1)
    style_heading(equipment_heading, level=1)
    
    kitchen_summary = data.get('equipment_inspection', [])
    
    if kitchen_summary:
        for kitchen_idx, kitchen in enumerate(kitchen_summary):
            # Kitchen header
            kitchen_title = doc.add_heading(f"Kitchen: {kitchen['name']}", level=2)
            style_heading(kitchen_title, level=2)
            
            # Process equipment in this kitchen
            for equip_idx, equip in enumerate(kitchen.get('equipment', [])):
                # Equipment header (as sub-section under kitchen)
                marvel_status = " (With Marvel)" if equip.get('with_marvel', False) else ""
                equip_title = doc.add_heading(f"  {equip['type_name']}{marvel_status}", level=3)
                style_heading(equip_title, level=3)
            
                # Equipment info
                equip_info_para = doc.add_paragraph()
                equip_info_para.add_run(f"Location: {equip['location']}").font.size = Pt(11)
                
                # Add alarm details if any alarms are registered
                if equip.get('alarm_details'):
                    doc.add_paragraph()
                    alarm_heading = doc.add_paragraph()
                    alarm_run = alarm_heading.add_run("Registered Alarms:")
                    alarm_run.bold = True
                    alarm_run.font.size = Pt(12)
                    alarm_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for alarms
                    
                    for alarm_key, alarm_data in equip['alarm_details'].items():
                        if alarm_data.get('description'):
                            alarm_num = alarm_key.replace('alarm_', '')
                            alarm_para = doc.add_paragraph()
                            alarm_para.add_run(f"Alarm {alarm_num}: ").bold = True
                            alarm_para.add_run(alarm_data['description']).font.size = Pt(11)
                            alarm_para.paragraph_format.left_indent = Inches(0.5)
                    
                    # Add alarm photos
                    alarm_photos = {}
                    for photo_key, photo_file in equip.get('photos', {}).items():
                        if 'photo_alarm_' in photo_key:
                            alarm_photos[photo_key] = photo_file
                    
                    if alarm_photos:
                        doc.add_paragraph()
                        alarm_photos_para = doc.add_paragraph()
                        alarm_photos_para.add_run("Alarm Photos:\n").bold = True
                        
                        # Group photos in pairs for side-by-side display
                        photo_items = list(alarm_photos.items())
                        for i in range(0, len(photo_items), 2):
                            # Create a table for side-by-side photos
                            photo_table = doc.add_table(rows=1, cols=2)
                            photo_table.autofit = False
                            
                            # First photo
                            photo_key, photo_file = photo_items[i]
                            cell1 = photo_table.cell(0, 0)
                            cell1_para = cell1.paragraphs[0]
                            cell1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Reset file position and add photo
                            photo_file.seek(0)
                            run1 = cell1_para.add_run()
                            run1.add_picture(photo_file, width=Inches(2.0))
                            
                            # Add caption
                            caption1 = cell1.add_paragraph()
                            caption_text = photo_key.replace('photo_', '').replace('_', ' ').title()
                            caption1.add_run(caption_text).font.size = Pt(9)
                            caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Second photo (if exists)
                            if i + 1 < len(photo_items):
                                photo_key2, photo_file2 = photo_items[i + 1]
                                cell2 = photo_table.cell(0, 1)
                                cell2_para = cell2.paragraphs[0]
                                cell2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                # Reset file position and add photo
                                photo_file2.seek(0)
                                run2 = cell2_para.add_run()
                                run2.add_picture(photo_file2, width=Inches(2.0))
                                
                                # Add caption
                                caption2 = cell2.add_paragraph()
                                caption_text2 = photo_key2.replace('photo_', '').replace('_', ' ').title()
                                caption2.add_run(caption_text2).font.size = Pt(9)
                                caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Add spacing after photos
                            doc.add_paragraph()
                
                # COMBINED INSPECTION FINDINGS TABLE
                all_findings = []
                
                # Add positive findings
                if equip.get('yes_responses'):
                    for yes_item in equip['yes_responses']:
                        question_text = yes_item.get('question', yes_item['item'].replace('_', ' ').title())
                        # Use the actual answer value (could be "Yes" or text like "GOT")
                        answer_text = yes_item.get('answer', 'YES')
                        if yes_item['comment']:
                            answer_text += f"\n{yes_item['comment']}"
                        all_findings.append((question_text, answer_text))
                
                # Add issues (NO responses)
                if equip.get('no_responses'):
                    for no_item in equip['no_responses']:
                        question_text = no_item.get('question', no_item['item'].replace('_', ' ').title())
                        # Use the actual answer value
                        answer_text = no_item.get('answer', 'NO')
                        if no_item['comment']:
                            answer_text += f"\n{no_item['comment']}"
                        all_findings.append((question_text, answer_text))
                
                # Add N/A responses
                if equip.get('na_responses'):
                    for na_item in equip['na_responses']:
                        question_text = na_item.get('question', na_item['item'].replace('_', ' ').title())
                        # Use the actual answer value
                        answer_text = na_item.get('answer', 'N/A')
                        if na_item['comment']:
                            answer_text += f"\n{na_item['comment']}"
                        all_findings.append((question_text, answer_text))
                
                # Create combined table if there are any findings
                if all_findings:
                    doc.add_paragraph()
                    findings_heading = doc.add_paragraph()
                    findings_run = findings_heading.add_run("Inspection Findings:")
                    findings_run.bold = True
                    findings_run.font.size = Pt(12)
                    
                    create_info_table(doc, all_findings, col_widths=[4, 2.5])
                
                # Combine all photos below the table
                all_photos = {}
                if equip.get('yes_photos'):
                    all_photos.update(equip['yes_photos'])
                if equip.get('no_photos'):
                    all_photos.update(equip['no_photos'])
                if equip.get('na_photos'):
                    all_photos.update(equip['na_photos'])
                
                # Add all photos if available
                if all_photos:
                    doc.add_paragraph()
                    photos_para = doc.add_paragraph()
                    photos_para.add_run("Supporting Photos:\n").bold = True
                    
                    # Group photos in pairs for side-by-side display
                    photo_items = list(all_photos.items())
                    for i in range(0, len(photo_items), 2):
                        # Create a table for side-by-side photos
                        photo_table = doc.add_table(rows=1, cols=2)
                        photo_table.autofit = False
                        
                        # First photo
                        photo_key, photo_file = photo_items[i]
                        cell1 = photo_table.cell(0, 0)
                        cell1_para = cell1.paragraphs[0]
                        cell1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Reset file position and add photo
                        photo_file.seek(0)
                        run1 = cell1_para.add_run()
                        run1.add_picture(photo_file, width=Inches(2.0))
                        
                        # Add caption
                        caption1 = cell1.add_paragraph()
                        caption1.add_run(photo_key.replace('photo_', '').replace('_', ' ').title()).font.size = Pt(9)
                        caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Second photo (if exists)
                        if i + 1 < len(photo_items):
                            photo_key2, photo_file2 = photo_items[i + 1]
                            cell2 = photo_table.cell(0, 1)
                            cell2_para = cell2.paragraphs[0]
                            cell2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Reset file position and add photo
                            photo_file2.seek(0)
                            run2 = cell2_para.add_run()
                            run2.add_picture(photo_file2, width=Inches(2.0))
                            
                            # Add caption
                            caption2 = cell2.add_paragraph()
                            caption2.add_run(photo_key2.replace('photo_', '').replace('_', ' ').title()).font.size = Pt(9)
                            caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add spacing after photos
                        doc.add_paragraph()
                
                # If no issues found at all
                if not equip.get('no_responses'):
                    doc.add_paragraph()
                    para = doc.add_paragraph()
                    no_issues_run = para.add_run("No issues identified during inspection.")
                    no_issues_run.font.size = Pt(11)
                    no_issues_run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                
                # Add spacing between equipment in the same kitchen
                if equip_idx < len(kitchen.get('equipment', [])) - 1:
                    doc.add_paragraph()
            
            # Add spacing between kitchens
            if kitchen_idx < len(kitchen_summary) - 1:
                doc.add_paragraph()
    else:
        para = doc.add_paragraph()
        para.add_run("No equipment inspection data available.").font.size = Pt(11)
    
    # WORK PERFORMED SECTION (only if filled)
    section_number = 3
    if data.get('work_performed'):
        work_heading = doc.add_heading(f'{section_number}. JOB DETAILS', level=1)
        style_heading(work_heading, level=1)
        
        work_para = doc.add_paragraph()
        work_text = work_para.add_run(data.get('work_performed', ''))
        work_text.font.size = Pt(11)
        work_para.paragraph_format.line_spacing = 1.5
        work_para.paragraph_format.space_after = Pt(12)
        section_number += 1
    
    # SPARE PARTS SECTION
    spare_parts = data.get('spare_parts', [])
    if spare_parts and any(part.get('name') for part in spare_parts):
        parts_heading = doc.add_heading(f'{section_number}. SPARE PARTS REQUIRED', level=1)
        style_heading(parts_heading, level=1)
        
        # Create table for spare parts
        parts_table = doc.add_table(rows=1, cols=3)
        parts_table.allow_autofit = False
        
        # Set column widths
        for cell in parts_table.columns[0].cells:
            cell.width = Inches(0.8)
        for cell in parts_table.columns[1].cells:
            cell.width = Inches(4.5)
        for cell in parts_table.columns[2].cells:
            cell.width = Inches(1.2)
        
        # Header row
        header_cells = parts_table.rows[0].cells
        header_cells[0].text = 'S.No.'
        header_cells[1].text = 'Spare Part Name'
        header_cells[2].text = 'Quantity'
        
        # Style header
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set background color
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '1f4788')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Add spare parts rows
        serial_no = 1
        for part in spare_parts:
            if part.get('name'):  # Only add if part name is not empty
                row = parts_table.add_row()
                row.cells[0].text = str(serial_no)
                row.cells[1].text = part.get('name', '')
                row.cells[2].text = str(part.get('quantity', 1))
                
                # Center align serial number and quantity
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                serial_no += 1
        
        # Apply professional table formatting
        format_table_style_enhanced(parts_table)
        
        # Add spacing after table
        doc.add_paragraph()
        section_number += 1
    else:
        # Add note if no spare parts required
        no_parts_heading = doc.add_heading(f'{section_number}. SPARE PARTS REQUIRED', level=1)
        style_heading(no_parts_heading, level=1)
        
        no_parts_para = doc.add_paragraph()
        no_parts_text = no_parts_para.add_run('No spare parts required.')
        no_parts_text.font.size = Pt(11)
        no_parts_text.font.italic = True
        no_parts_para.paragraph_format.space_after = Pt(12)
        section_number += 1
    
    # RECOMMENDATIONS SECTION
    if data.get('recommendations'):
        rec_heading = doc.add_heading(f'{section_number}. RECOMMENDATIONS', level=1)
        style_heading(rec_heading, level=1)
        
        rec_para = doc.add_paragraph()
        rec_text = rec_para.add_run(data.get('recommendations', ''))
        rec_text.font.size = Pt(11)
        rec_para.paragraph_format.line_spacing = 1.5
        rec_para.paragraph_format.space_after = Pt(12)
        section_number += 1
    
    # Add page break before signatures
    doc.add_page_break()
    
    # SIGNATURE SECTION
    sig_heading = doc.add_heading('ACKNOWLEDGMENT AND SIGNATURES', level=1)
    style_heading(sig_heading, level=1)
    
    # Add acknowledgment text
    ack_para = doc.add_paragraph()
    ack_text = ack_para.add_run(
        "The undersigned acknowledge that the service described in this report has been "
        "completed satisfactorily and in accordance with the agreed specifications."
    )
    ack_text.font.size = Pt(10)
    ack_text.font.italic = True
    ack_para.paragraph_format.space_after = Pt(24)
    
    # Create signature table
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Configure signature columns
    for col in sig_table.columns:
        for cell in col.cells:
            cell.width = Inches(3)
    
    # Technician signature
    sig_table.cell(0, 0).text = "Service Technician:"
    
    # Add signature image if available
    if data.get('technician_signature'):
        sig_cell = sig_table.cell(1, 0)
        sig_para = sig_cell.paragraphs[0]
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sig_para.add_run()
        run.add_picture(data.get('technician_signature'), width=Inches(1.5))
    else:
        sig_table.cell(1, 0).text = "_" * 35
    
    sig_table.cell(2, 0).text = data.get('technician_name', '')
    sig_table.cell(3, 0).text = f"Date: {datetime.now().strftime('%B %d, %Y')}"
    
    # Customer signature
    sig_table.cell(0, 1).text = "Customer Representative:"
    
    # Add customer signature image if available
    if data.get('customer_signature'):
        sig_cell = sig_table.cell(1, 1)
        sig_para = sig_cell.paragraphs[0]
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sig_para.add_run()
        run.add_picture(data.get('customer_signature'), width=Inches(1.5))
    else:
        sig_table.cell(1, 1).text = "_" * 35
    
    sig_table.cell(2, 1).text = data.get('customer_signatory', data.get('customer_name', ''))
    sig_table.cell(3, 1).text = f"Date: {datetime.now().strftime('%B %d, %Y')}"
    
    # Style signature table
    for row in sig_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(11)
                run.font.name = 'Arial'
            set_cell_margins(cell, top=0.1, bottom=0.1)
    
    # Make labels bold
    sig_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    sig_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    
    # Add final note
    doc.add_paragraph()
    note_para = doc.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_text = note_para.add_run(
        "This report is confidential and proprietary.\n"
        "For service inquiries, please contact our Service Department."
    )
    note_text.font.size = Pt(9)
    note_text.font.name = 'Arial'
    note_text.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes


def create_general_service_report(data):
    """Generate a Professional General Service Report Word document"""
    
    # Helper function to set cell background
    def set_cell_background(cell, color):
        """Set background color for a table cell"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)
    
    # Helper function to set table borders
    def set_table_borders(table):
        """Apply borders to table"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Remove existing borders
        for child in tblPr:
            if child.tag.endswith('tblBorders'):
                tblPr.remove(child)
        
        # Create new borders
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    # Use template if available, otherwise create new document
    template_path = "Templates/Report Letter Head.docx"
    
    try:
        doc = Document(template_path)
    except:
        doc = Document()
    
    # Get the default style and set font
    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)
    except:
        pass
    
    # Set default font for headings
    try:
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'Arial'
    except:
        pass
    
    # Title
    title = doc.add_heading('GENERAL SERVICE REPORT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_heading(title, level=0)
    
    # Add minimal spacing
    doc.add_paragraph()
    
    # GENERAL INFORMATION SECTION
    general_heading = doc.add_heading('1. GENERAL INFORMATION', level=1)
    style_heading(general_heading, level=1)
    
    # Create professional info table (same style as Technical Report)
    general_info = [
        ("Customer Name", data.get('customer_name', '')),
        ("Project Name", data.get('project_name', '')),
        ("Contact Person", data.get('contact_person', '')),
        ("Location", data.get('outlet_location', '')),
        ("Contact Number", data.get('contact_number', '')),
        ("Visit Type", data.get('visit_type', '')),
        ("Visit Classification", data.get('visit_class', ''))
    ]
    
    create_info_table(doc, general_info)
    doc.add_paragraph()  # Add spacing
    
    # WORK PERFORMED SECTION
    section_number = 2
    work_heading = doc.add_heading(f'{section_number}. WORK PERFORMED', level=1)
    style_heading(work_heading, level=1)
    
    work_performed_list = data.get('work_performed_list', [])
    
    if work_performed_list:
        # Create a table for work performed items
        work_table = doc.add_table(rows=1, cols=2)
        work_table.allow_autofit = False
        work_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        work_table.columns[0].width = Inches(0.8)
        work_table.columns[1].width = Inches(5.7)
        
        # Header row
        header_cells = work_table.rows[0].cells
        header_cells[0].text = 'S.No.'
        header_cells[1].text = 'Work Performed Description'
        
        # Style header row to match Technical Report
        for i, cell in enumerate(header_cells):
            if i == 0:  # S.No. column
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:  # Description column
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
            set_cell_background(cell, 'E8E8E8')
        
        # Add work items
        for idx, work_item in enumerate(work_performed_list):
            if work_item.get('title') or work_item.get('description'):
                row = work_table.add_row()
                row.cells[0].text = str(idx + 1)
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Use title if available, otherwise use description
                row.cells[1].text = work_item.get('title', work_item.get('description', ''))
                
                # Apply font size and name
                for cell in row.cells:
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'
        
        # Apply borders
        set_table_borders(work_table)
        doc.add_paragraph()
        
        # Add detailed descriptions if any
        has_descriptions = any(work_item.get('description') for work_item in work_performed_list)
        if has_descriptions:
            details_heading = doc.add_paragraph()
            details_run = details_heading.add_run("Work Details:")
            details_run.bold = True
            details_run.font.size = Pt(12)
            details_run.font.name = 'Arial'
            
            for idx, work_item in enumerate(work_performed_list):
                if work_item.get('description'):
                    # Work item number and title
                    item_para = doc.add_paragraph()
                    item_title = item_para.add_run(f"{idx + 1}. {work_item.get('title', f'Work Item {idx + 1}')}: ")
                    item_title.bold = True
                    item_title.font.size = Pt(11)
                    item_title.font.name = 'Arial'
                    
                    # Description
                    desc_text = item_para.add_run(work_item['description'])
                    desc_text.font.size = Pt(11)
                    desc_text.font.name = 'Arial'
                    item_para.paragraph_format.left_indent = Inches(0.25)
                    item_para.paragraph_format.space_after = Pt(6)
            
            doc.add_paragraph()
        
        # Add photos section if any work items have photos
        has_photos = any(work_item.get('photos') for work_item in work_performed_list)
        if has_photos:
            photos_heading = doc.add_paragraph()
            photos_run = photos_heading.add_run("Work Photos:")
            photos_run.bold = True
            photos_run.font.size = Pt(12)
            photos_run.font.name = 'Arial'
            
            # Collect all photos from all work items
            all_photos = []
            for work_idx, work_item in enumerate(work_performed_list):
                if work_item.get('photos'):
                    for photo_idx, photo in enumerate(work_item['photos']):
                        photo_desc = work_item.get('photo_descriptions', {}).get(str(photo_idx), f'Work Item {work_idx + 1} - Photo {photo_idx + 1}')
                        all_photos.append((photo, photo_desc))
            
            # Display photos in pairs
            for i in range(0, len(all_photos), 2):
                photo_table = doc.add_table(rows=1, cols=2)
                photo_table.autofit = False
                
                # First photo
                photo_file, photo_desc = all_photos[i]
                cell1 = photo_table.cell(0, 0)
                cell1_para = cell1.paragraphs[0]
                cell1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Reset file position and add photo
                photo_file.seek(0)
                run1 = cell1_para.add_run()
                run1.add_picture(photo_file, width=Inches(2.0))
                
                # Add caption
                caption1 = cell1.add_paragraph()
                caption_run1 = caption1.add_run(photo_desc)
                caption_run1.font.size = Pt(9)
                caption_run1.font.name = 'Arial'
                caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Second photo (if exists)
                if i + 1 < len(all_photos):
                    photo_file2, photo_desc2 = all_photos[i + 1]
                    cell2 = photo_table.cell(0, 1)
                    cell2_para = cell2.paragraphs[0]
                    cell2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Reset file position and add photo
                    photo_file2.seek(0)
                    run2 = cell2_para.add_run()
                    run2.add_picture(photo_file2, width=Inches(2.0))
                    
                    # Add caption
                    caption2 = cell2.add_paragraph()
                    caption_run2 = caption2.add_run(photo_desc2)
                    caption_run2.font.size = Pt(9)
                    caption_run2.font.name = 'Arial'
                    caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add spacing after photos
                doc.add_paragraph()
    else:
        no_work_para = doc.add_paragraph()
        no_work_text = no_work_para.add_run('No work performed recorded.')
        no_work_text.font.size = Pt(11)
        no_work_text.font.name = 'Arial'
        no_work_text.font.italic = True
        no_work_para.paragraph_format.space_after = Pt(12)
    
    section_number += 1
    
    # SPARE PARTS SECTION (same as technical report)
    spare_parts = data.get('spare_parts', [])
    if spare_parts and any(part.get('name') for part in spare_parts):
        parts_heading = doc.add_heading(f'{section_number}. SPARE PARTS REQUIRED', level=1)
        style_heading(parts_heading, level=1)
        
        # Create table for spare parts
        parts_table = doc.add_table(rows=1, cols=3)
        parts_table.allow_autofit = False
        parts_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        parts_table.columns[0].width = Inches(0.8)
        parts_table.columns[1].width = Inches(4.5)
        parts_table.columns[2].width = Inches(1.2)
        
        # Header row
        header_cells = parts_table.rows[0].cells
        header_cells[0].text = 'S.No.'
        header_cells[1].text = 'Spare Part Description'
        header_cells[2].text = 'Quantity'
        
        # Style header row to match Technical Report
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
            set_cell_background(cell, 'E8E8E8')
        
        # Add spare parts
        for idx, part in enumerate(spare_parts):
            if part.get('name'):
                row = parts_table.add_row()
                row.cells[0].text = str(idx + 1)
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[1].text = part.get('name', '')
                row.cells[2].text = str(part.get('quantity', 1))
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Apply font size and name
                for cell in row.cells:
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'
        
        # Apply borders
        set_table_borders(parts_table)
        doc.add_paragraph()
        section_number += 1
    else:
        parts_heading = doc.add_heading(f'{section_number}. SPARE PARTS REQUIRED', level=1)
        style_heading(parts_heading, level=1)
        
        no_parts_para = doc.add_paragraph()
        no_parts_text = no_parts_para.add_run('No spare parts required.')
        no_parts_text.font.size = Pt(11)
        no_parts_text.font.name = 'Arial'
        no_parts_text.font.italic = True
        no_parts_para.paragraph_format.space_after = Pt(12)
        section_number += 1
    
    # RECOMMENDATIONS SECTION (same as technical report)
    if data.get('recommendations'):
        rec_heading = doc.add_heading(f'{section_number}. RECOMMENDATIONS', level=1)
        style_heading(rec_heading, level=1)
        
        rec_para = doc.add_paragraph()
        rec_text = rec_para.add_run(data.get('recommendations', ''))
        rec_text.font.size = Pt(11)
        rec_text.font.name = 'Arial'
        rec_para.paragraph_format.line_spacing = 1.5
        rec_para.paragraph_format.space_after = Pt(12)
        section_number += 1
    
    # Add page break before signatures
    doc.add_page_break()
    
    # SIGNATURE SECTION - Updated to match Testing & Commissioning format
    sig_heading = doc.add_heading('ACKNOWLEDGMENT AND SIGNATURES', level=1)
    style_heading(sig_heading, level=1)
    
    # Add acknowledgment text
    ack_para = doc.add_paragraph()
    ack_text = ack_para.add_run(
        "The undersigned acknowledge that the service described in this report has been "
        "completed satisfactorily and in accordance with the agreed specifications."
    )
    ack_text.font.size = Pt(11)
    ack_text.font.name = 'Arial'
    ack_text.font.italic = True
    ack_para.paragraph_format.space_after = Pt(36)
    
    # Add some space before the signature area
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Create a 2-column table for side-by-side signatures
    sig_table = doc.add_table(rows=5, cols=2)
    sig_table.allow_autofit = False
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for cell in sig_table.columns[0].cells:
        cell.width = Inches(3.2)
    for cell in sig_table.columns[1].cells:
        cell.width = Inches(3.2)
    
    # Service Technician section
    tech_label = sig_table.cell(0, 0)
    tech_label.text = "Service Technician:"
    tech_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in tech_label.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
    
    # Add space for signature
    sig_space = sig_table.cell(1, 0)
    sig_space_para = sig_space.paragraphs[0]
    sig_space_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if data.get('technician_signature'):
        run = sig_space_para.add_run()
        run.add_picture(data.get('technician_signature'), width=Inches(2))
    else:
        # Create empty space for signature
        sig_space_para.add_run("\n\n\n\n")
    
    # Signature line
    sig_line = sig_table.cell(2, 0)
    sig_line.text = "_" * 30
    sig_line.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Empty row for spacing
    sig_table.cell(3, 0).text = ""
    
    # Date
    date_cell = sig_table.cell(4, 0)
    date_cell.text = f"Date: {datetime.now().strftime('%B %d, %Y')}"
    date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Customer Representative section
    cust_label = sig_table.cell(0, 1)
    cust_label.text = "Customer Representative:"
    cust_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cust_label.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
    
    # Add space for customer signature
    cust_space = sig_table.cell(1, 1)
    cust_space_para = cust_space.paragraphs[0]
    cust_space_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if data.get('customer_signature'):
        run = cust_space_para.add_run()
        run.add_picture(data.get('customer_signature'), width=Inches(2))
    else:
        # Create empty space for signature
        cust_space_para.add_run("\n\n\n\n")
    
    # Signature line
    cust_line = sig_table.cell(2, 1)
    cust_line.text = "_" * 30
    cust_line.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Empty row for spacing
    sig_table.cell(3, 1).text = ""
    
    # Date
    cust_date = sig_table.cell(4, 1)
    cust_date.text = f"Date: {datetime.now().strftime('%B %d, %Y')}"
    cust_date.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Style all text in signature table
    for row in sig_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if not run.font.bold:
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'
            set_cell_margins(cell, top=0.1, bottom=0.1)
    
    # Add footer text
    doc.add_paragraph()
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = footer_para.add_run(
        "This report is confidential and proprietary.\n"
        "For service inquiries, please contact our Service Department."
    )
    footer_text.font.size = Pt(9)
    footer_text.font.name = 'Arial'
    footer_text.font.color.rgb = RGBColor(128, 128, 128)
    footer_text.font.italic = True
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

def format_tc_table(table, is_header=False):
    """Format Testing & Commissioning table with blue header style and reduced row height"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Add table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Format cells and set row heights
    for row_idx, row in enumerate(table.rows):
        # Set row height
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '280')  # Further reduced height in twips (280 = ~0.19 inches)
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        
        for cell in row.cells:
            # Set minimal cell margins
            set_cell_margins(cell, top=0.02, bottom=0.02, left=0.08, right=0.08)
            
            # Add vertical alignment (middle)
            tcPr = cell._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
            
            # Style the cell based on position
            if is_header and row_idx == 0:
                # Blue header background
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), '1F4788')  # Blue color
                tcPr.append(shading)
                
                # White bold text for header
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Reduce paragraph spacing
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
            else:
                # Regular cells - center all text horizontally
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Reduce paragraph spacing
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Arial'

def create_testing_commissioning_report(data):
    """Generate a Testing and Commissioning Report Word document"""
    import math
    
    # Use template if available, otherwise create new document
    template_path = "Templates/Report Letter Head.docx"
    
    try:
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
                section.header_distance = Inches(0.5)
                section.footer_distance = Inches(0.5)
    except Exception as e:
        st.warning(f"Could not load template: {str(e)}. Using blank document.")
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)
    
    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('TESTING AND COMMISSIONING REPORT')
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(31, 71, 136)
    title_run.font.bold = True
    
    # Add report date
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref_para.add_run(f"Report Date: {data.get('date', datetime.now().strftime('%B %d, %Y'))}")
    ref_run.font.size = Pt(11)
    ref_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # GENERAL INFORMATION SECTION
    general_heading = doc.add_heading('GENERAL INFORMATION', level=1)
    style_heading(general_heading, level=1)
    
    # Create professional info table
    general_info = [
        ("Customer Name", data.get('customer_name', '')),
        ("Project Name", data.get('project_name', '')),
        ("Contact Person", data.get('contact_person', '')),
        ("Location", data.get('outlet_location', '')),
        ("Contact Number", data.get('contact_number', '')),
        ("Visit Type", data.get('visit_type', '')),
        ("Visit Classification", data.get('visit_class', ''))
    ]
    
    create_info_table(doc, general_info)
    
    # Add page break before canopy commissioning data
    doc.add_page_break()
    
    # CANOPY COMMISSIONING DATA SECTION
    canopy_heading = doc.add_heading('CANOPY COMMISSIONING DATA', level=1)
    style_heading(canopy_heading, level=1)
    
    canopy_data = data.get('canopy_data', [])
    
    for canopy_idx, canopy in enumerate(canopy_data):
        # Add page break before each canopy (except the first one since we already have a page break)
        if canopy_idx > 0:
            doc.add_page_break()
        
        # Check if this is a Mobichef model (only has checklist, no air data)
        if canopy.get('model') == 'Mobichef':
            # For Mobichef, create a MOBICHEF DATA table
            # Create table with header row for "MOBICHEF DATA"
            mobichef_header_table = doc.add_table(rows=1, cols=1)
            mobichef_header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            mobichef_header_cell = mobichef_header_table.cell(0, 0)
            mobichef_header_cell.text = "MOBICHEF DATA"
            format_tc_table(mobichef_header_table, is_header=True)
            
            # Mobichef Info Table (merged with header visually)
            mobichef_info_table = doc.add_table(rows=3, cols=2)
            mobichef_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Populate mobichef info (only relevant fields)
            mobichef_info_table.cell(0, 0).text = "Drawing Number"
            mobichef_info_table.cell(0, 1).text = canopy.get('drawing_number', '')
            mobichef_info_table.cell(1, 0).text = "Canopy Location"
            mobichef_info_table.cell(1, 1).text = canopy.get('location', '')
            mobichef_info_table.cell(2, 0).text = "Canopy Model"
            mobichef_info_table.cell(2, 1).text = canopy.get('model', '')
            
            format_tc_table(mobichef_info_table)
            
            doc.add_paragraph()
        else:
            # EXTRACT AIR DATA for non-Mobichef models
            # Create table with header row for "EXTRACT AIR DATA"
            extract_header_table = doc.add_table(rows=1, cols=1)
            extract_header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            extract_header_cell = extract_header_table.cell(0, 0)
            extract_header_cell.text = "EXTRACT AIR DATA"
            format_tc_table(extract_header_table, is_header=True)
            
            # Extract Air Info Table (merged with header visually)
            extract_info_table = doc.add_table(rows=6, cols=2)
            extract_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Populate extract info
            extract_info_table.cell(0, 0).text = "Drawing Number"
            extract_info_table.cell(0, 1).text = canopy.get('drawing_number', '')
            extract_info_table.cell(1, 0).text = "Canopy Location"
            extract_info_table.cell(1, 1).text = canopy.get('location', '')
            extract_info_table.cell(2, 0).text = "Canopy Model"
            extract_info_table.cell(2, 1).text = canopy.get('model', '')
            # Get extract data first
            extract_data = canopy.get('extract_data', [])
            # Sum design flowrates from all modules (convert from L/s to m³/s)
            total_extract_design_ls = sum(module.get('design_flowrate_ls', 0.0) for module in extract_data)
            total_extract_design_m3s = total_extract_design_ls / 1000  # Convert L/s to m³/s
            extract_info_table.cell(3, 0).text = "Design Flowrate"
            extract_info_table.cell(3, 1).text = f"{total_extract_design_m3s:.2f} m³/s"
            extract_info_table.cell(4, 0).text = "Quantity of Canopy Sections"
            extract_info_table.cell(4, 1).text = str(canopy.get('modules', 1))
            extract_info_table.cell(5, 0).text = "Calculation"
            # Show correct formula based on hood type
            is_cmw_type = canopy.get('model') in ['CMWF', 'CMWI', 'CMW-MUAP-CJ', 'CMW-CJ', 'CMW', 'CXW']
            if is_cmw_type:
                extract_info_table.cell(5, 1).text = "QE = V × L × W"
            else:
                extract_info_table.cell(5, 1).text = "Qv = K √Pa"
            
            format_tc_table(extract_info_table)
            
            # Extract Air Readings Table (connected to info table)
            extract_data = canopy.get('extract_data', [])
            if extract_data:
                # Check if CMW type to determine columns
                is_cmw = canopy.get('model') in ['CMWF', 'CMWI', 'CMW-MUAP-CJ', 'CMW-CJ']
                
                if is_cmw:
                    # CMW type table with different columns
                    extract_table = doc.add_table(rows=len(extract_data) + 2, cols=7)  # +2 for header and total
                    extract_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Headers for CMW
                    headers = ['Hood #', 'Anemometer Reading\n(V - m/s)', 'Length of\nopening\n(mm)', 
                              'Width of\nopening\n(meter)', 'Achieved\n(m³/h)', 'Design\n(L/s)', 'Percentage']
                    for col_idx, header in enumerate(headers):
                        cell = extract_table.cell(0, col_idx)
                        cell.text = header
                    
                    total_achieved = 0
                    total_design = 0
                    
                    # Data rows for CMW
                    for row_idx, section_data in enumerate(extract_data):
                        achieved_m3s = section_data.get('flowrate_m3s', 0.0)
                        achieved_m3h = achieved_m3s * 3600  # Convert to m³/h
                        
                        # Get design flowrate for this module (stored in L/s, convert to m³/h for display)
                        design_ls = section_data.get('design_flowrate_ls', 0.0)
                        design_m3s = design_ls / 1000  # Convert L/s to m³/s
                        design_m3h = design_m3s * 3600  # Convert to m³/h
                        
                        extract_table.cell(row_idx + 1, 0).text = f"M{row_idx + 1}"
                        extract_table.cell(row_idx + 1, 1).text = f"{section_data.get('anemometer', 0.0):.2f} m/s"
                        # Display length in mm (stored in mm, displayed in mm)
                        length_mm = section_data.get('length_opening', 1800)
                        extract_table.cell(row_idx + 1, 2).text = f"{length_mm:.0f}"
                        extract_table.cell(row_idx + 1, 3).text = "0.09m"
                        extract_table.cell(row_idx + 1, 4).text = f"{achieved_m3h:.2f}"  # Display in m³/h
                        extract_table.cell(row_idx + 1, 5).text = f"{design_ls:.0f}"  # Display in L/s
                        extract_table.cell(row_idx + 1, 6).text = f"{section_data.get('percentage', 0):.0f}%"
                        
                        total_achieved += achieved_m3h
                        total_design += design_ls  # Sum in L/s
                    
                    # Total row - only populate and border columns 3-6
                    # Leave first 3 columns empty (no borders)
                    total_row_idx = len(extract_data) + 1
                    
                    # Remove borders from first 3 cells of total row
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    
                    for col_idx in range(3):
                        cell = extract_table.cell(total_row_idx, col_idx)
                        tcPr = cell._tc.get_or_add_tcPr()
                        # Remove all borders for these cells
                        tcBorders = OxmlElement('w:tcBorders')
                        for border_name in ['top', 'left', 'bottom', 'right']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set(qn('w:val'), 'nil')
                            tcBorders.append(border)
                        tcPr.append(tcBorders)
                    
                    # Style the TOTAL cell with blue background and white text
                    total_cell = extract_table.cell(total_row_idx, 3)
                    total_cell.text = "TOTAL"
                    
                    # Apply blue background to TOTAL cell
                    total_tcPr = total_cell._tc.get_or_add_tcPr()
                    total_shading = OxmlElement('w:shd')
                    total_shading.set(qn('w:fill'), '1F4788')  # Blue color
                    total_tcPr.append(total_shading)
                    
                    # Make text white and bold
                    for paragraph in total_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(10)
                            run.font.name = 'Arial'
                    
                    extract_table.cell(total_row_idx, 4).text = f"{total_achieved:.0f}"
                    extract_table.cell(total_row_idx, 5).text = f"{total_design:.0f}"
                    # Convert total_design from L/s to m³/h for percentage calculation
                    if total_design > 0:
                        total_design_m3h = (total_design / 1000) * 3600  # L/s to m³/s to m³/h
                        total_percentage = (total_achieved / total_design_m3h) * 100
                    else:
                        total_percentage = 0
                    extract_table.cell(total_row_idx, 6).text = f"{total_percentage:.0f}%"
                else:
                    # Regular table with K-Factor
                    extract_table = doc.add_table(rows=len(extract_data) + 2, cols=7)  # +2 for header and total
                    extract_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Headers
                    headers = ['Module', 'Manometer\nReading (Pa)', 'K-Factor\n(m³/h)', 'Flowrate\n(m³/h)', 'Flowrate\n(m³/s)', 'Design\n(L/s)', 'Percentage']
                    for col_idx, header in enumerate(headers):
                        cell = extract_table.cell(0, col_idx)
                        cell.text = header
                    
                    total_achieved_m3h = 0
                    total_achieved_m3s = 0
                    total_design_ls = 0
                    
                    # Data rows
                    for row_idx, section_data in enumerate(extract_data):
                        flowrate_m3h = section_data.get('flowrate_m3h', 0)
                        flowrate_m3s = section_data.get('flowrate_m3s', 0.0)
                        design_ls = section_data.get('design_flowrate_ls', 0.0)
                        
                        extract_table.cell(row_idx + 1, 0).text = f"M{row_idx + 1}"
                        extract_table.cell(row_idx + 1, 1).text = f"{section_data.get('tab_reading', 0.0):.1f}"
                        extract_table.cell(row_idx + 1, 2).text = f"{section_data.get('k_factor', 0.0):.1f}"
                        extract_table.cell(row_idx + 1, 3).text = f"{flowrate_m3h:.0f}"
                        extract_table.cell(row_idx + 1, 4).text = f"{flowrate_m3s:.3f}"
                        extract_table.cell(row_idx + 1, 5).text = f"{design_ls:.0f}"
                        extract_table.cell(row_idx + 1, 6).text = f"{section_data.get('percentage', 0):.0f}%"
                        
                        total_achieved_m3h += flowrate_m3h
                        total_achieved_m3s += flowrate_m3s
                        total_design_ls += design_ls
                    
                    # Total row
                    total_row_idx = len(extract_data) + 1
                    
                    # Add borders to total row starting from column 2 (index 2)
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    
                    # First two columns empty (no borders)
                    for col_idx in range(2):
                        cell = extract_table.cell(total_row_idx, col_idx)
                        tcPr = cell._tc.get_or_add_tcPr()
                        tcBorders = OxmlElement('w:tcBorders')
                        for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right']:
                            border = OxmlElement(border_name)
                            border.set(qn('w:val'), 'nil')
                            tcBorders.append(border)
                        tcPr.append(tcBorders)
                    
                    # Column 2: "TOTAL" text
                    total_cell = extract_table.cell(total_row_idx, 2)
                    total_cell.text = "TOTAL"
                    total_tcPr = total_cell._tc.get_or_add_tcPr()
                    total_shading = OxmlElement('w:shd')
                    total_shading.set(qn('w:val'), 'clear')
                    total_shading.set(qn('w:color'), 'auto')
                    total_shading.set(qn('w:fill'), '1F4788')  # Blue color
                    total_tcPr.append(total_shading)
                    
                    # Make text white and bold
                    for paragraph in total_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(10)
                            run.font.name = 'Arial'
                    
                    # Fill in total values
                    extract_table.cell(total_row_idx, 3).text = f"{total_achieved_m3h:.0f}"
                    extract_table.cell(total_row_idx, 4).text = f"{total_achieved_m3s:.3f}"
                    extract_table.cell(total_row_idx, 5).text = f"{total_design_ls:.0f}"
                    
                    # Calculate percentage
                    if total_design_ls > 0:
                        # Convert design from L/s to m³/h for percentage calculation
                        total_design_m3h = (total_design_ls / 1000) * 3600
                        total_percentage = (total_achieved_m3h / total_design_m3h) * 100
                    else:
                        total_percentage = 0
                    extract_table.cell(total_row_idx, 6).text = f"{total_percentage:.0f}%"
                
                format_tc_table(extract_table, is_header=True)
            
            doc.add_paragraph()
            
            # SUPPLY AIR DATA (if applicable)
            if canopy.get('model') in ['KVF', 'UVF', 'CMWF', 'CMW-MUAP-CJ', 'KVD', 'KVV']:
                # SUPPLY AIR DATA
                # Create table with header row for "SUPPLY AIR DATA"
                supply_header_table = doc.add_table(rows=1, cols=1)
                supply_header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                supply_header_cell = supply_header_table.cell(0, 0)
                supply_header_cell.text = "SUPPLY AIR DATA"
                format_tc_table(supply_header_table, is_header=True)
                
                # Supply Air Info Table (merged with header visually)
                supply_info_table = doc.add_table(rows=6, cols=2)
                supply_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Populate supply info
                supply_info_table.cell(0, 0).text = "Drawing Number"
                supply_info_table.cell(0, 1).text = canopy.get('drawing_number', '')
                supply_info_table.cell(1, 0).text = "Canopy Location"
                supply_info_table.cell(1, 1).text = canopy.get('location', '')
                supply_info_table.cell(2, 0).text = "Canopy Model"
                supply_info_table.cell(2, 1).text = canopy.get('model', '')
                # Get supply data first
                supply_data = canopy.get('supply_data', [])
                # Sum design flowrates from all modules (convert from L/s to m³/s)
                total_supply_design_ls = sum(module.get('design_flowrate_ls', 0.0) for module in supply_data)
                total_supply_design_m3s = total_supply_design_ls / 1000  # Convert L/s to m³/s
                supply_info_table.cell(3, 0).text = "Design Flowrate"
                supply_info_table.cell(3, 1).text = f"{total_supply_design_m3s:.2f} m³/s"
                supply_info_table.cell(4, 0).text = "Quantity of Canopy Sections"
                supply_info_table.cell(4, 1).text = str(canopy.get('modules', 1))
                supply_info_table.cell(5, 0).text = "Calculation"
                # Supply always uses K-Factor calculation
                supply_info_table.cell(5, 1).text = "Qv = K √Pa"
                
                format_tc_table(supply_info_table)
                
                # Supply Air Readings Table (connected to info table)
                if supply_data:
                    supply_table = doc.add_table(rows=len(supply_data) + 1, cols=7)
                    supply_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Headers
                    headers = ['Module', 'Manometer\nReading (Pa)', 'K-Factor\n(m³/h)', 'Flowrate\n(m³/h)', 'Flowrate\n(m³/s)', 'Design\n(L/s)', 'Percentage']
                    for col_idx, header in enumerate(headers):
                        cell = supply_table.cell(0, col_idx)
                        cell.text = header
                    
                    # Data rows
                    for row_idx, section_data in enumerate(supply_data):
                        supply_table.cell(row_idx + 1, 0).text = f"M{row_idx + 1}"
                        supply_table.cell(row_idx + 1, 1).text = f"{section_data.get('tab_reading', 0.0):.1f}"
                        supply_table.cell(row_idx + 1, 2).text = f"{section_data.get('k_factor', 0.0):.1f}"
                        supply_table.cell(row_idx + 1, 3).text = f"{section_data.get('flowrate_m3h', 0):.0f}"
                        supply_table.cell(row_idx + 1, 4).text = f"{section_data.get('flowrate_m3s', 0.0):.3f}"
                        supply_table.cell(row_idx + 1, 5).text = f"{section_data.get('design_flowrate_ls', 0.0):.0f}"
                        supply_table.cell(row_idx + 1, 6).text = f"{section_data.get('percentage', 0):.0f}%"
                    
                    format_tc_table(supply_table, is_header=True)
        
        
        # EQUIPMENT CHECKLIST for this canopy
        # Get checklist for this specific canopy using location and model
        canopy_location = canopy.get('location', f'Canopy {canopy_idx+1}')
        canopy_model = canopy.get('model', 'Unknown')
        checklist_key = f'{canopy_location} {canopy_model}'
        
        checklists_data = data.get('tc_checklists', {})
        if checklist_key in checklists_data and checklists_data[checklist_key]:
            checklist_items = checklists_data[checklist_key]
            
            # Add space before checklist
            doc.add_paragraph()
            
            # Create header table for checklist
            checklist_header_table = doc.add_table(rows=1, cols=1)
            checklist_header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            checklist_header_cell = checklist_header_table.cell(0, 0)
            checklist_header_cell.text = f"{canopy_location} {canopy_model} EQUIPMENT CHECKLIST"
            format_tc_table(checklist_header_table, is_header=True)
            
            # Create checklist items table (no headers)
            checklist_table = doc.add_table(rows=len(checklist_items), cols=2)
            checklist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Checklist items
            for row_idx, (item_name, status) in enumerate(checklist_items.items()):
                checklist_table.cell(row_idx, 0).text = item_name
                # Handle different status values
                if status == "Yes" or status == "OK" or status == "Clean":
                    display_status = "✓"
                elif status == "No" or status == "Faulty" or status == "Dirty" or status == "Overload" or status == "Missing":
                    display_status = "✗"
                elif status == "N/A":
                    display_status = "N/A"
                else:
                    display_status = status  # Display as-is for any other status
                checklist_table.cell(row_idx, 1).text = display_status
            
            format_tc_table(checklist_table, is_header=False)  # No blue header for checklist items
        
        doc.add_paragraph()
    
    # RECOMMENDATIONS SECTION
    rec_heading = doc.add_heading('RECOMMENDATIONS', level=1)
    style_heading(rec_heading, level=1)
    
    recommendations_para = doc.add_paragraph()
    recommendations_para.add_run(data.get('recommendations', 'No specific recommendations at this time.'))
    
    # Add page break before signatures
    doc.add_page_break()
    
    # SIGNATURES SECTION
    sig_heading = doc.add_heading('SIGNATURES', level=1)
    style_heading(sig_heading, level=1)
    
    # Create signature table
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.allow_autofit = False
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for cell in sig_table.columns[0].cells:
        cell.width = Inches(3)
    for cell in sig_table.columns[1].cells:
        cell.width = Inches(3)
    
    # Technician signature
    sig_table.cell(0, 0).text = "Technician:"
    
    # Add technician signature image if available
    if data.get('technician_signature'):
        sig_cell = sig_table.cell(1, 0)
        sig_para = sig_cell.paragraphs[0]
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sig_para.add_run()
        run.add_picture(data.get('technician_signature'), width=Inches(1.5))
    else:
        sig_table.cell(1, 0).text = "_" * 35
    
    sig_table.cell(2, 0).text = data.get('technician_name', '')
    sig_table.cell(3, 0).text = f"Date: {datetime.now().strftime('%B %d, %Y')}"
    
    # Customer signature
    sig_table.cell(0, 1).text = "Customer Representative:"
    sig_table.cell(1, 1).text = "_" * 35
    sig_table.cell(2, 1).text = data.get('customer_name', '')
    sig_table.cell(3, 1).text = f"Date: {datetime.now().strftime('%B %d, %Y')}"
    
    # Style signature table
    for row in sig_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(11)
                run.font.name = 'Arial'
            set_cell_margins(cell, top=0.1, bottom=0.1)
    
    # Make labels bold
    sig_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    sig_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    
    # Add footer with Halton branding if not using template
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = footer_para.add_run(
        "Halton Middle East LTD\n"
        "Zip Code 23218, 2163 Building 7571\n"
        "C.R: 4030531945\n"
        "VAT: 311906248400003\n"
        "Tel. +966 (012) 6614490"
    )
    footer_text.font.size = Pt(9)
    footer_text.font.name = 'Arial'
    footer_text.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

def main():
    # Check for shared form data in URL parameters
    query_params = st.experimental_get_query_params()
    
    # Debug: Show what parameters we have
    if query_params:
        st.sidebar.write("🔍 URL Parameters detected:", dict(query_params))
    
    if 'data' in query_params and not st.session_state.get('data_restored', False):
        st.info("🔄 Restoring form data from shared link...")
        encoded_data = query_params['data']
        
        # Debug: Show encoded data length
        st.sidebar.write(f"📊 Encoded data length: {len(encoded_data)} characters")
        
        decoded_data = decode_form_data_from_url(encoded_data)
        if decoded_data:
            # Debug: Show what data we decoded
            st.sidebar.write("✅ Data decoded successfully")
            st.sidebar.write(f"Customer: {decoded_data.get('basic_info', {}).get('customer_name', 'N/A')}")
            st.sidebar.write(f"Kitchens: {decoded_data.get('kitchen_data', {}).get('num_kitchens', 0)}")
            
            if restore_form_data(decoded_data):
                st.success("✅ Form data restored from shared link!")
                st.session_state['data_restored'] = True
                st.rerun()
            else:
                st.error("❌ Failed to restore form data from link")
        else:
            st.error("❌ Failed to decode shared link data")
    
    # Header
    st.markdown('<h1 class="main-header">Service Reports System</h1>', unsafe_allow_html=True)
    
    # Sidebar for report type selection
    with st.sidebar:
        st.markdown("### Report Type Selection")
        report_type = st.selectbox(
            "Select Report Type",
            ["Technical Report", "General Service Report", "Testing and Commissioning Report"],
            index=0
        )
        
        # Remove the coming soon check for Testing and Commissioning Report
    
    # Main form based on report type
    if report_type == "Technical Report":
        st.markdown('<h2 class="section-header">Technical Report Form</h2>', unsafe_allow_html=True)
    elif report_type == "Testing and Commissioning Report":
        st.markdown('<h2 class="section-header">Testing and Commissioning Report</h2>', unsafe_allow_html=True)
    else:  # General Service Report
        st.markdown('<h2 class="section-header">General Service Report Form</h2>', unsafe_allow_html=True)
    
    # General Information Section (outside form)
    st.markdown("### General Information")
    col1, col2 = st.columns(2)
    
    with col1:
        customer_name = st.text_input("Customer's Name*", placeholder="e.g., ABC Company", key="customer_name")
        project_name = st.text_input("Project Name*", placeholder="e.g., Kitchen equipment service", key="project_name")
        contact_person = st.text_input("Contact Person*", placeholder="e.g., John Smith", key="contact_person")
        outlet_location = st.text_input("Outlet/Location*", placeholder="e.g., Downtown Location", key="outlet_location")
    
    with col2:
        contact_number = st.text_input("Contact #*", placeholder="e.g., +1 555 123 4567", key="contact_number")
        
        # Set default visit type based on report type
        visit_type_options = ["", "Service Call", "AMC (Contract)", "Emergency Service", "Installation", "Commissioning", "Inspection (Healthy Visit)"]
        if report_type == "Testing and Commissioning Report":
            # Set Commissioning as default for Testing and Commissioning Report
            default_index = visit_type_options.index("Commissioning")
        else:
            default_index = 0
        
        visit_type = st.selectbox(
            "Visit Type*",
            visit_type_options,
            index=default_index,
            key="visit_type"
        )
        
        visit_class = st.selectbox(
            "Visit Class*",
            ["To Be Invoiced", "Free of Charge", "Warranty"],
            key="visit_class"
        )
        date = st.date_input("Report Date", value=datetime.now(), key="report_date")
    
    # Testing & Commissioning Report specific sections
    if report_type == "Testing and Commissioning Report":
        # Import canopy models and configuration
        CANOPY_MODELS = ["", "KVF", "KVI", "UVF", "CMW", "CXW", "CMWF", "CMWI", "CMW-MUAP-CJ", "CMW-CJ", "KVD", "KVV", "Mobichef"]
        
        st.markdown("### Canopy Configuration")
        
        # Initialize session state for canopies
        if 'canopy_data' not in st.session_state:
            st.session_state.canopy_data = []
        
        # Add first canopy if list is empty
        if len(st.session_state.canopy_data) == 0:
            st.session_state.canopy_data.append({
                'drawing_number': '',
                'location': '',
                'model': '',
                'modules': 1,
                'extract_data': [],
                'supply_data': []
            })
        
        # Function to add a new hood
        def add_new_hood():
            st.session_state.canopy_data.append({
                'drawing_number': '',
                'location': '',
                'model': '',
                'modules': 1,
                'extract_data': [],
                'supply_data': []
            })
        
        # Display canopy configuration forms
        for i in range(len(st.session_state.canopy_data)):
            with st.expander(f"🏭 Hood {i+1}", expanded=True):
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.text_input(
                        "Drawing Number",
                        key=f"drawing_{i}",
                        value=st.session_state.canopy_data[i].get('drawing_number', ''),
                        on_change=lambda idx=i: st.session_state.canopy_data[idx].update({'drawing_number': st.session_state[f"drawing_{idx}"]})
                    )
                
                with col2:
                    st.text_input(
                        "Canopy Location",
                        key=f"location_{i}",
                        value=st.session_state.canopy_data[i].get('location', ''),
                        on_change=lambda idx=i: st.session_state.canopy_data[idx].update({'location': st.session_state[f"location_{idx}"]})
                    )
                
                with col3:
                    # Get current model for this canopy
                    current_model = st.session_state.canopy_data[i].get('model', '')
                    st.selectbox(
                        "Canopy Model",
                        options=CANOPY_MODELS,
                        key=f"model_{i}",
                        index=CANOPY_MODELS.index(current_model) if current_model in CANOPY_MODELS else 0,
                        placeholder="Select a model",
                        on_change=lambda idx=i: st.session_state.canopy_data[idx].update({'model': st.session_state[f"model_{idx}"]})
                    )
                    # Always read from session state key to ensure we have the latest value
                    canopy_model = st.session_state.get(f"model_{i}", st.session_state.canopy_data[i].get('model', ''))
                
                # Only show Module Configuration and other sections if a model is selected
                if canopy_model and canopy_model != '':
                    # Module Configuration (not for Mobichef)
                    if canopy_model != 'Mobichef':
                        st.markdown("#### Module Configuration")
                        
                        st.number_input(
                            "Quantity of Canopy Modules",
                            min_value=1,
                            max_value=10,
                            key=f"modules_{i}",
                            value=st.session_state.canopy_data[i].get('modules', 1),
                            on_change=lambda idx=i: st.session_state.canopy_data[idx].update({'modules': st.session_state[f"modules_{idx}"]})
                        )
                        num_modules = st.session_state.canopy_data[i].get('modules', 1)
                    else:
                        # Mobichef has no modules
                        num_modules = 1
                        st.session_state.canopy_data[i]['modules'] = 1
                    
                    # Check if model has supply (F models or CMW-MUAP-CJ) - use the most current value
                    current_canopy_model = st.session_state.get(f"model_{i}", st.session_state.canopy_data[i].get('model', ''))
                    # Models with supply: all F models plus CMW-MUAP-CJ
                    has_supply = ('F' in current_canopy_model or current_canopy_model == 'CMW-MUAP-CJ') if current_canopy_model else False
                    
                    # Initialize module data
                    while len(st.session_state.canopy_data[i]['extract_data']) < num_modules:
                        st.session_state.canopy_data[i]['extract_data'].append({
                            'num_ksa_filters': 1,
                            'tab_reading': 0.0,
                            'k_factor': 0.0,
                            'flowrate_m3h': 0.0,
                            'flowrate_m3s': 0.0,
                            'design_flowrate_ls': 0.0  # Design flowrate in L/s
                        })
                        st.session_state.canopy_data[i]['supply_data'].append({
                            'hood_length': 1.0,
                            'tab_reading': 0.0,
                            'k_factor': 0.0,
                            'flowrate_m3h': 0.0,
                            'flowrate_m3s': 0.0,
                            'design_flowrate_ls': 0.0  # Design flowrate in L/s
                        })
                    
                    # Trim module data if needed
                    st.session_state.canopy_data[i]['extract_data'] = st.session_state.canopy_data[i]['extract_data'][:num_modules]
                    st.session_state.canopy_data[i]['supply_data'] = st.session_state.canopy_data[i]['supply_data'][:num_modules]
                    
                    # Extract Air Data (not applicable for Mobichef)
                    if canopy_model != 'Mobichef':
                        st.markdown("#### Extract Air Data")
                    
                    # Determine if model has UV (UVF/UVI models)
                    has_uv = canopy_model in ['UVF', 'UVI'] if canopy_model else False
                    # Check if it's a CMW type (water wash hood)
                    is_cmw = canopy_model in ['CMWF', 'CMWI', 'CMW-MUAP-CJ', 'CMW-CJ'] if canopy_model else False
                    
                    for j in range(num_modules):
                        st.markdown(f"**Module {j+1}**")
                        
                        if is_cmw:
                            # CMW type uses different calculation: QE = V × L × W
                            col1, col2, col3, col4, col5, col6, col7 = st.columns(7)
                        
                        if is_cmw:
                            # CMW type calculation
                            with col1:
                                st.number_input(
                                    "Anemometer (m/s)",
                                    min_value=0.0,
                                    step=0.1,
                                    key=f"anemometer_{i}_{j}",
                                    value=st.session_state.canopy_data[i]['extract_data'][j].get('anemometer', 0.0),
                                    help="Anemometer reading in m/s",
                                    on_change=lambda idx=i, jdx=j: st.session_state.canopy_data[idx]['extract_data'][jdx].update({'anemometer': st.session_state[f"anemometer_{idx}_{jdx}"]})
                                )
                            
                            with col2:
                                st.number_input(
                                    "Length (mm)",
                                    min_value=0,
                                    step=100,
                                    key=f"length_opening_{i}_{j}",
                                    value=int(st.session_state.canopy_data[i]['extract_data'][j].get('length_opening', 1800)),
                                    help="Length of opening in millimeters",
                                    on_change=lambda idx=i, jdx=j: st.session_state.canopy_data[idx]['extract_data'][jdx].update({'length_opening': st.session_state[f"length_opening_{idx}_{jdx}"]})
                                )
                            
                            with col3:
                                width_opening = st.number_input(
                                    "Width (m)",
                                    value=0.09,
                                    disabled=True,
                                    key=f"width_opening_{i}_{j}",
                                    help="Width of opening (fixed at 0.09m)"
                                )
                                st.session_state.canopy_data[i]['extract_data'][j]['width_opening'] = 0.09
                            
                            with col4:
                                # Design flowrate input in L/s
                                st.number_input(
                                    "Design (L/s)",
                                    min_value=0.0,
                                    step=10.0,
                                    key=f"extract_design_{i}_{j}",
                                    value=st.session_state.canopy_data[i]['extract_data'][j].get('design_flowrate_ls', 0.0),
                                    on_change=lambda idx=i, jdx=j: st.session_state.canopy_data[idx]['extract_data'][jdx].update({'design_flowrate_ls': st.session_state[f"extract_design_{idx}_{jdx}"]})
                                )
                            
                            # Calculate flowrate using QE = V × L × W
                            # Get the most current values
                            current_anemometer = st.session_state.canopy_data[i]['extract_data'][j].get('anemometer', 0.0)
                            current_length_mm = st.session_state.canopy_data[i]['extract_data'][j].get('length_opening', 1800)
                            current_length_m = current_length_mm / 1000  # Convert mm to m
                            flowrate_m3s = current_anemometer * current_length_m * width_opening
                            flowrate_m3h = flowrate_m3s * 3600
                            st.session_state.canopy_data[i]['extract_data'][j]['flowrate_m3h'] = flowrate_m3h
                            st.session_state.canopy_data[i]['extract_data'][j]['flowrate_m3s'] = flowrate_m3s
                            
                            with col5:
                                st.text_input(
                                    "Flowrate (m³/h)",
                                    value=f"{flowrate_m3h:.0f}",
                                    disabled=True,
                                    key=f"extract_m3h_{i}_{j}"
                                )
                            
                            with col6:
                                st.text_input(
                                    "Flowrate (m³/s)",
                                    value=f"{flowrate_m3s:.3f}",
                                    disabled=True,
                                    key=f"extract_m3s_{i}_{j}"
                                )
                            
                            # Calculate percentage using module's design flowrate
                            design_flowrate_ls = st.session_state.canopy_data[i]['extract_data'][j].get('design_flowrate_ls', 0.0)
                            if design_flowrate_ls > 0:
                                # Convert L/s to m³/s for comparison
                                design_flowrate_m3s = design_flowrate_ls / 1000
                                percentage = (flowrate_m3s / design_flowrate_m3s) * 100
                            else:
                                percentage = 0
                            st.session_state.canopy_data[i]['extract_data'][j]['percentage'] = percentage
                            
                            with col7:
                                st.text_input(
                                    "Percentage",
                                    value=f"{percentage:.0f}%",
                                    disabled=True,
                                    key=f"extract_percentage_{i}_{j}"
                                )
                        else:
                            # Regular calculation for non-CMW types
                            col1, col2, col3, col4, col5, col6, col7 = st.columns(7)
                            
                            with col1:
                                # Number of KSA filters for this module
                                st.number_input(
                                    "KSA Filters",
                                    min_value=1,
                                    max_value=6,
                                    key=f"ksa_filters_{i}_{j}",
                                    value=st.session_state.canopy_data[i]['extract_data'][j].get('num_ksa_filters', 1),
                                    help="Number of KSA filters for this module",
                                    on_change=lambda idx=i, jdx=j: st.session_state.canopy_data[idx]['extract_data'][jdx].update({'num_ksa_filters': st.session_state[f"ksa_filters_{idx}_{jdx}"]})
                                )
                                num_ksa_filters = st.session_state.canopy_data[i]['extract_data'][j].get('num_ksa_filters', 1)
                            
                            with col2:
                                # Design flowrate input in L/s
                                st.number_input(
                                    "Design (L/s)",
                                    min_value=0.0,
                                    step=10.0,
                                    key=f"extract_design_{i}_{j}",
                                    value=st.session_state.canopy_data[i]['extract_data'][j].get('design_flowrate_ls', 0.0),
                                    on_change=lambda idx=i, jdx=j: st.session_state.canopy_data[idx]['extract_data'][jdx].update({'design_flowrate_ls': st.session_state[f"extract_design_{idx}_{jdx}"]})
                                )
                            
                            with col3:
                                st.number_input(
                                    "Manometer Reading (Pa)",
                                    min_value=0.0,
                                    step=0.1,
                                    key=f"extract_tab_{i}_{j}",
                                    value=st.session_state.canopy_data[i]['extract_data'][j].get('tab_reading', 0.0),
                                    format="%.1f",
                                    help="Enter the manometer reading in Pascals",
                                    on_change=lambda idx=i, jdx=j: st.session_state.canopy_data[idx]['extract_data'][jdx].update({'tab_reading': st.session_state[f"extract_tab_{idx}_{jdx}"]})
                                )
                            
                            # Get automatic K-Factor based on this module's number of filters
                            # Use the most current value
                            current_ksa = st.session_state.canopy_data[i]['extract_data'][j].get('num_ksa_filters', 1)
                            auto_k_factor_extract = get_extract_k_factor(current_ksa, canopy_model, has_uv)
                            
                            with col4:
                                # Display auto-calculated K-Factor (read-only)
                                st.text_input(
                                    "K-Factor (m³/h)",
                                    value=f"{auto_k_factor_extract:.1f}",
                                    disabled=True,
                                    key=f"extract_k_display_{i}_{j}",
                                    help=f"Auto-calculated based on {num_ksa_filters} KSA filter(s)"
                                )
                                # Store the K-Factor in data
                                st.session_state.canopy_data[i]['extract_data'][j]['k_factor'] = auto_k_factor_extract
                            
                            # Calculate flowrates
                            # Use the most current values
                            current_tab = st.session_state.canopy_data[i]['extract_data'][j].get('tab_reading', 0.0)
                            if current_tab > 0 and auto_k_factor_extract > 0:
                                import math
                                flowrate_m3h = auto_k_factor_extract * math.sqrt(current_tab)
                                flowrate_m3s = flowrate_m3h / 3600
                                st.session_state.canopy_data[i]['extract_data'][j]['flowrate_m3h'] = flowrate_m3h
                                st.session_state.canopy_data[i]['extract_data'][j]['flowrate_m3s'] = flowrate_m3s
                            else:
                                flowrate_m3h = 0
                                flowrate_m3s = 0
                            
                            with col5:
                                st.text_input(
                                    "Flowrate (m³/h)",
                                    value=f"{flowrate_m3h:.0f}",
                                    disabled=True,
                                    key=f"extract_m3h_{i}_{j}"
                                )
                            
                            with col6:
                                st.text_input(
                                    "Flowrate (m³/s)",
                                    value=f"{flowrate_m3s:.3f}",
                                    disabled=True,
                                    key=f"extract_m3s_{i}_{j}"
                                )
                            
                            # Calculate percentage using module's design flowrate
                            design_flowrate_ls = st.session_state.canopy_data[i]['extract_data'][j].get('design_flowrate_ls', 0.0)
                            if design_flowrate_ls > 0:
                                # Convert L/s to m³/s for comparison
                                design_flowrate_m3s = design_flowrate_ls / 1000
                                percentage = (flowrate_m3s / design_flowrate_m3s) * 100
                            else:
                                percentage = 0
                            st.session_state.canopy_data[i]['extract_data'][j]['percentage'] = percentage
                            
                            with col7:
                                st.text_input(
                                    "Percentage",
                                    value=f"{percentage:.0f}%",
                                    disabled=True,
                                    key=f"extract_percentage_{i}_{j}"
                                )
                
                # Supply Air Data (if model has supply)
                if canopy_model in ['KVF', 'UVF', 'CMWF', 'CMW-MUAP-CJ', 'KVD', 'KVV']:
                    st.markdown("#### Supply Air Data")
                    
                    for j in range(num_modules):
                        st.markdown(f"**Module {j+1}**")
                        col1, col2, col3, col4, col5, col6, col7 = st.columns(7)
                        
                        with col1:
                            # Hood/Plenum length for this module
                            st.number_input(
                                "Hood Length (m)",
                                min_value=1.0,
                                max_value=4.0,
                                step=0.1,
                                format="%.1f",
                                key=f"hood_length_{i}_{j}",
                                value=st.session_state.canopy_data[i]['supply_data'][j].get('hood_length', 1.0),
                                help="Hood length for this module",
                                on_change=lambda idx=i, jdx=j: st.session_state.canopy_data[idx]['supply_data'][jdx].update({'hood_length': st.session_state[f"hood_length_{idx}_{jdx}"]})
                            )
                            hood_length = st.session_state.canopy_data[i]['supply_data'][j].get('hood_length', 1.0)
                        
                        with col2:
                            # Design flowrate input in L/s
                            st.number_input(
                                "Design (L/s)",
                                min_value=0.0,
                                step=10.0,
                                key=f"supply_design_{i}_{j}",
                                value=st.session_state.canopy_data[i]['supply_data'][j].get('design_flowrate_ls', 0.0),
                                on_change=lambda idx=i, jdx=j: st.session_state.canopy_data[idx]['supply_data'][jdx].update({'design_flowrate_ls': st.session_state[f"supply_design_{idx}_{jdx}"]})
                            )
                        
                        with col3:
                            st.number_input(
                                "Manometer Reading (Pa)",
                                min_value=0.0,
                                step=0.1,
                                key=f"supply_tab_{i}_{j}",
                                value=st.session_state.canopy_data[i]['supply_data'][j].get('tab_reading', 0.0),
                                format="%.1f",
                                help="Enter the manometer reading in Pascals",
                                on_change=lambda idx=i, jdx=j: st.session_state.canopy_data[idx]['supply_data'][jdx].update({'tab_reading': st.session_state[f"supply_tab_{idx}_{jdx}"]})
                            )
                        
                        # Get automatic K-Factor based on this module's hood length
                        current_hood_length = st.session_state.canopy_data[i]['supply_data'][j].get('hood_length', 1.0)
                        auto_k_factor_supply = get_supply_k_factor(current_hood_length)
                        
                        with col4:
                            # Display auto-calculated K-Factor (read-only)
                            st.text_input(
                                "K-Factor (m³/h)",
                                value=f"{auto_k_factor_supply:.1f}",
                                disabled=True,
                                key=f"supply_k_display_{i}_{j}",
                                help=f"Auto-calculated based on {hood_length:.1f}m hood length"
                            )
                            # Store the K-Factor in data
                            st.session_state.canopy_data[i]['supply_data'][j]['k_factor'] = auto_k_factor_supply
                        
                        # Calculate flowrates
                        current_supply_tab = st.session_state.canopy_data[i]['supply_data'][j].get('tab_reading', 0.0)
                        if current_supply_tab > 0 and auto_k_factor_supply > 0:
                            import math
                            flowrate_m3h = auto_k_factor_supply * math.sqrt(current_supply_tab)
                            flowrate_m3s = flowrate_m3h / 3600
                            st.session_state.canopy_data[i]['supply_data'][j]['flowrate_m3h'] = flowrate_m3h
                            st.session_state.canopy_data[i]['supply_data'][j]['flowrate_m3s'] = flowrate_m3s
                        else:
                            flowrate_m3h = 0
                            flowrate_m3s = 0
                        
                        with col5:
                            st.text_input(
                                "Flowrate (m³/h)",
                                value=f"{flowrate_m3h:.0f}",
                                disabled=True,
                                key=f"supply_m3h_{i}_{j}"
                            )
                        
                        with col6:
                            st.text_input(
                                "Flowrate (m³/s)",
                                value=f"{flowrate_m3s:.3f}",
                                disabled=True,
                                key=f"supply_m3s_{i}_{j}"
                            )
                        
                        # Calculate percentage using module's design flowrate
                        design_flowrate_ls = st.session_state.canopy_data[i]['supply_data'][j].get('design_flowrate_ls', 0.0)
                        if design_flowrate_ls > 0:
                            # Convert L/s to m³/s for comparison
                            design_flowrate_m3s = design_flowrate_ls / 1000
                            percentage = (flowrate_m3s / design_flowrate_m3s) * 100
                        else:
                            percentage = 0
                        st.session_state.canopy_data[i]['supply_data'][j]['percentage'] = percentage
                        
                        with col7:
                            st.text_input(
                                "Percentage",
                                value=f"{percentage:.0f}%",
                                disabled=True,
                                key=f"supply_percentage_{i}_{j}"
                            )
                
                # Inspection Checklist (only for Testing & Commissioning Report)
                if report_type == "Testing and Commissioning Report":
                    st.markdown("---")  # Add separator
                    canopy_location = st.session_state.canopy_data[i].get('location', f'Canopy {i+1}')
                    canopy_model = st.session_state.canopy_data[i].get('model', 'Unknown')
                    st.markdown(f"#### {canopy_location} {canopy_model} Equipment Checklist")
            
                    # Initialize session state for this canopy's checklist if not exists
                    if 'tc_checklists' not in st.session_state:
                        st.session_state.tc_checklists = {}
                    
                    # Create a more descriptive key using location and model
                    canopy_location = st.session_state.canopy_data[i].get('location', f'Canopy {i+1}')
                    canopy_model = st.session_state.canopy_data[i].get('model', 'Unknown')
                    checklist_key = f'{canopy_location} {canopy_model}'
                    
                    if checklist_key not in st.session_state.tc_checklists:
                        st.session_state.tc_checklists[checklist_key] = {}
                
                    # Define checklist items based on model type
                    checklist_items = []
                
                    if canopy_model in ['KVI', 'KVF', 'KVD', 'KVV']:
                        checklist_items = [
                            "All filters are in place",
                            "Hood lights are working",
                            "Capture Jet Fan is Working"
                        ]
                    elif canopy_model in ['UVF', 'UVI']:
                        checklist_items = [
                            "Communication to all UV Controllers confirmed",
                            "Safety interlock to UV Door/chambers tested for correct operation", 
                            "Safety interlock to all filters tested for correct operation",
                            "UV operation tested and left working",
                            "HOOD lights are working",
                            "CAPTURE JET FAN is working"
                        ]
                    
                        # Check if water wash system exists
                        has_water_wash = st.checkbox(
                            "UV Hood has Water Wash System",
                            key=f"uv_water_wash_{i}",
                            value=st.session_state.canopy_data[i].get('has_water_wash', False)
                        )
                        st.session_state.canopy_data[i]['has_water_wash'] = has_water_wash
                    
                        if has_water_wash:
                            checklist_items.extend([
                                "Checked operation of Water Nozzles",
                                "Checked Hot Water pressure (2 Bar at 65 Degrees)",
                                "Checked Hot Water flow to the Hoods",
                                "Checked the detergent pump operation",
                                "Checked the detergent chemical flow to the hood"
                            ])
                        
                    elif canopy_model in ['CMWF', 'CMWI', 'CMW', 'CXW', 'CMW-MUAP-CJ', 'CMW-CJ']:
                        checklist_items = [
                            "Checked operation of Water Nozzles",
                            "Checked Cold Water pressure (2 BAR)",
                            "Checked Hot Water pressure (2 BAR)",
                            "Checked Water flow to the Hoods",
                            "Checked the detergent pump operation",
                            "Checked the detergent chemical flow to the hood",
                            "HOOD lights are working",
                            "CAPTURE JET FAN is working"
                        ]
                    elif canopy_model == 'Mobichef':
                        checklist_items = [
                            "Hood lights status",
                            "All KSA Filters in place",
                            "All Mesh Filters in place",
                            "Check pre filter status",
                            "Check ESP status",
                            "Check Carbon Filter status",
                            "Check Carbon Particle filter status"
                        ]
                
                    # Add Marvel system checklist if needed
                    has_marvel = st.checkbox(
                        "Include Marvel System Checklist",
                        key=f"marvel_{i}",
                        value=st.session_state.canopy_data[i].get('has_marvel', False)
                    )
                    st.session_state.canopy_data[i]['has_marvel'] = has_marvel
                
                    if has_marvel:
                        marvel_items = [
                            "Display Touch Screen tested for correct operation",
                            "Control Calculator power up and tested for correct operation",
                            "Settings configured (0-10Vdc / 4-20mA)",
                            "Exhaust Hood ABD dampers are operational",
                            "Infrared Sensor (IR Sensor) tested for operation and Direction Adjusted",
                            "Room Sensor installed and tested",
                            "No Alarm on the Touch Screen"
                        ]
                        checklist_items.extend(marvel_items)
                
                    # Display checklist items with dropdowns
                    if checklist_items:
                        # Display main checklist header
                        st.markdown("##### Main Equipment Checklist")
                        
                        # Track if headers have been shown
                        water_wash_header_shown = False
                        marvel_header_shown = False
                        
                        # Determine where water wash items start (for UV models with water wash)
                        water_wash_start_item = "Checked operation of Water Nozzles"
                        
                        # Determine where Marvel items start
                        marvel_start_item = "Display Touch Screen tested for correct operation"
                        
                        for idx, item in enumerate(checklist_items):
                            # Add section headers when needed (only once)
                            if item == water_wash_start_item and not water_wash_header_shown:
                                st.markdown("##### Water Wash System Checklist")
                                water_wash_header_shown = True
                            elif item == marvel_start_item and not marvel_header_shown:
                                st.markdown("##### Marvel System Checklist")
                                marvel_header_shown = True
                            
                            col1, col2 = st.columns([3, 1])
                            with col1:
                                st.write(item)
                            with col2:
                                # Different dropdown options for Mobichef
                                if canopy_model == 'Mobichef':
                                    # Determine options based on item type
                                    if 'lights' in item.lower():
                                        options = ["OK", "Faulty", "N/A"]
                                        default_index = 0  # Default to "OK"
                                    elif 'All KSA Filters in place' in item or 'All Mesh Filters in place' in item:
                                        # Special case for KSA and Mesh filters
                                        options = ["OK", "Missing", "N/A"]
                                        default_index = 0  # Default to "OK"
                                    elif 'filter' in item.lower() or 'esp' in item.lower():
                                        options = ["Clean", "Dirty", "Overload", "Missing", "N/A"]
                                        default_index = 0  # Default to "Clean"
                                    else:
                                        options = ["OK", "Faulty", "N/A"]
                                        default_index = 0  # Default to "OK"
                                    
                                    status = st.selectbox(
                                        "",
                                        options=options,
                                        key=f"checklist_{i}_{idx}",
                                        label_visibility="collapsed",
                                        index=default_index
                                    )
                                else:
                                    # Yes/No for other models
                                    status = st.selectbox(
                                        "",
                                        options=["Yes", "No"],
                                        key=f"checklist_{i}_{idx}",
                                        label_visibility="collapsed",
                                        index=0  # Default to "Yes" (first option)
                                    )
                                # Store in session state
                            st.session_state.tc_checklists[checklist_key][item] = status
                else:
                    # No model selected - ensure basic initialization
                    # Make sure extract_data and supply_data are initialized as empty lists
                    if 'extract_data' not in st.session_state.canopy_data[i]:
                        st.session_state.canopy_data[i]['extract_data'] = []
                    if 'supply_data' not in st.session_state.canopy_data[i]:
                        st.session_state.canopy_data[i]['supply_data'] = []
            
            # Delete button at the bottom of each expander (only show if more than one hood)
            if len(st.session_state.canopy_data) > 1:
                # Get location and model for delete button text
                delete_location = st.session_state.canopy_data[i].get('location', '')
                delete_model = st.session_state.canopy_data[i].get('model', '')
                delete_text = f"🗑️ Delete"
                if delete_location:
                    delete_text += f" {delete_location}"
                if delete_model:
                    delete_text += f" {delete_model}"
                if not delete_location and not delete_model:
                    delete_text += f" Hood {i+1}"
                
                if st.button(delete_text, key=f"delete_hood_{i}", use_container_width=True):
                    st.session_state.canopy_data.pop(i)
                    # Clean up checklist data for this hood
                    if i < len(st.session_state.canopy_data):
                        canopy_location = st.session_state.canopy_data[i].get('location', f'Hood {i+1}')
                        canopy_model = st.session_state.canopy_data[i].get('model', 'Unknown')
                    else:
                        canopy_location = f'Hood {i+1}'
                        canopy_model = 'Unknown'
                    checklist_key = f'{canopy_location} {canopy_model}'
                    if 'tc_checklists' in st.session_state and checklist_key in st.session_state.tc_checklists:
                        del st.session_state.tc_checklists[checklist_key]
                    st.rerun()
        
        # Single Add Hood button at the bottom
        if st.button("➕ Add Hood", key="add_hood_main", use_container_width=True):
            add_new_hood()
            st.rerun()
    
    # Kitchen and Equipment Inspection Section (only for Technical Report)
    elif report_type == "Technical Report":
        st.markdown("### Kitchen and Equipment Inspection")
        
        # Number of kitchens
        # Initialize session state for kitchen count if not exists
        if "num_kitchens" not in st.session_state:
            st.session_state["num_kitchens"] = len(st.session_state.kitchen_list) if st.session_state.kitchen_list else 1
    
        num_kitchens = st.number_input(
            "Number of Kitchens",
            min_value=1,
            key="num_kitchens"
        )
        
        # Get the number from session state to ensure consistency
        num_kitchens = st.session_state["num_kitchens"]
        
        # Initialize or adjust kitchen list
        if len(st.session_state.kitchen_list) != num_kitchens:
            if len(st.session_state.kitchen_list) < num_kitchens:
                # Add new kitchens
                for i in range(len(st.session_state.kitchen_list), num_kitchens):
                    kitchen_id = f"kitchen_{i}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                    st.session_state.kitchen_list.append({
                        'id': kitchen_id,
                        'name': f'Kitchen {i + 1}',
                        'equipment_list': []
                    })
            else:
                # Remove kitchens
                st.session_state.kitchen_list = st.session_state.kitchen_list[:num_kitchens]
        
        # Display kitchens and their equipment
        for kitchen_idx, kitchen in enumerate(st.session_state.kitchen_list):
            with st.expander(f"🏪 Kitchen #{kitchen_idx + 1}", expanded=True):
                # Kitchen name input
                kitchen_name_key = f"kitchen_name_{kitchen_idx}"
                if kitchen_name_key not in st.session_state:
                    st.session_state[kitchen_name_key] = kitchen.get('name', f'Kitchen {kitchen_idx + 1}')
                
                kitchen_name = st.text_input(
                    "Kitchen Name",
                    key=kitchen_name_key,
                    placeholder="e.g., Main Kitchen, Prep Kitchen, etc."
                )
                kitchen['name'] = st.session_state[kitchen_name_key]
                
                # Number of equipment in this kitchen
                num_equipment_key = f"num_equipment_{kitchen_idx}"
                
                # Initialize session state for equipment count if not exists
                if num_equipment_key not in st.session_state:
                    st.session_state[num_equipment_key] = len(kitchen.get('equipment_list', []))
                
                num_equipment = st.number_input(
                    f"Number of Equipment in {kitchen['name']}",
                    min_value=0,
                    max_value=20,
                    key=num_equipment_key
                )
                
                # Get the number from session state to ensure consistency
                num_equipment = st.session_state[num_equipment_key]
                
                # Initialize or adjust equipment list for this kitchen
                if 'equipment_list' not in kitchen:
                    kitchen['equipment_list'] = []
                    
                current_count = len(kitchen['equipment_list'])
                if current_count != num_equipment:
                    if current_count < num_equipment:
                        # Add new equipment - only add what's needed
                        items_to_add = num_equipment - current_count
                        for i in range(items_to_add):
                            equipment_id = f"equipment_{kitchen_idx}_{current_count + i}_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
                            kitchen['equipment_list'].append({
                                'id': equipment_id,
                                'type': '',
                                'with_marvel': False,
                                'location': '',
                                'inspection_data': {},
                                'photos': {}
                            })
                    elif current_count > num_equipment:
                        # Remove excess equipment
                        kitchen['equipment_list'] = kitchen['equipment_list'][:num_equipment]
                
                # Display equipment for this kitchen
                if kitchen['equipment_list']:
                    st.markdown(f"#### Equipment in {kitchen['name']} ({len(kitchen['equipment_list'])} items)")
                    for equip_idx, equipment in enumerate(kitchen['equipment_list']):
                        with st.container():
                            st.markdown(f"**Equipment #{equip_idx + 1}**")
                            col1, col2 = st.columns(2)
                            
                            # Create unique keys including kitchen index
                            equipment_key_prefix = f"k{kitchen_idx}_e{equip_idx}"
                            
                            with col1:
                                # Equipment type selection
                                equip_type_key = f"equip_type_{equipment_key_prefix}"
                                # Initialize if not exists
                                if equip_type_key not in st.session_state:
                                    st.session_state[equip_type_key] = equipment.get('type', '')
                                
                                # Store previous equipment type to detect changes
                                previous_equip_type = equipment.get('type', '')
                                    
                                st.selectbox(
                                    "Equipment Type*",
                                    options=[''] + list(EQUIPMENT_TYPES.keys()),
                                    format_func=lambda x: EQUIPMENT_TYPES[x]["name"] if x else "Select equipment type",
                                    key=equip_type_key
                                )
                                # Update the equipment data from session state
                                equipment['type'] = st.session_state[equip_type_key]
                                
                                # If equipment type changed, clear all inspection data for this equipment
                                if previous_equip_type != equipment['type'] and previous_equip_type:
                                    # Clear all inspection data (except Marvel-prefixed ones)
                                    if 'inspection_data' in equipment:
                                        keys_to_remove = []
                                        for key in equipment['inspection_data'].keys():
                                            if not key.startswith('marvel_'):  # Keep Marvel data if exists
                                                keys_to_remove.append(key)
                                        for key in keys_to_remove:
                                            del equipment['inspection_data'][key]
                                    
                                    # Clear all photos (except Marvel-related ones)
                                    if 'photos' in equipment:
                                        keys_to_remove = []
                                        for key in equipment['photos'].keys():
                                            if 'marvel_' not in key:  # Keep Marvel photos if exists
                                                keys_to_remove.append(key)
                                        for key in keys_to_remove:
                                            del equipment['photos'][key]
                                    
                                    # Clear related session state keys for this equipment
                                    keys_to_clear = []
                                    for key in st.session_state.keys():
                                        if equipment_key_prefix in key and (key.startswith('q_') or key.startswith('comment_') or key.startswith('photo_')):
                                            if 'marvel_' not in key:  # Don't clear Marvel-related keys
                                                keys_to_clear.append(key)
                                    for key in keys_to_clear:
                                        del st.session_state[key]
                                
                                # Only show Marvel checkbox if equipment type is not ECOLOGY
                                if equipment['type'] != 'ECOLOGY':
                                    marvel_key = f"with_marvel_{equipment_key_prefix}"
                                    # Initialize if not exists
                                    if marvel_key not in st.session_state:
                                        st.session_state[marvel_key] = equipment.get('with_marvel', False)
                                        
                                    # Store previous state to detect changes
                                    previous_marvel_state = equipment.get('with_marvel', False)
                                    
                                    with_marvel = st.checkbox(
                                        "With Marvel System",
                                        key=marvel_key
                                    )
                                    # Update the equipment data from session state
                                    equipment['with_marvel'] = st.session_state[marvel_key]
                                else:
                                    # For ECOLOGY, always set with_marvel to False
                                    equipment['with_marvel'] = False
                                    previous_marvel_state = False
                                
                                # If Marvel was unchecked, clear all Marvel-related data
                                if previous_marvel_state and not equipment['with_marvel']:
                                    # Clear Marvel inspection data
                                    if 'inspection_data' in equipment:
                                        keys_to_remove = []
                                        for key in equipment['inspection_data'].keys():
                                            if key.startswith('marvel_'):
                                                keys_to_remove.append(key)
                                        for key in keys_to_remove:
                                            del equipment['inspection_data'][key]
                                    
                                    # Clear Marvel photos
                                    if 'photos' in equipment:
                                        keys_to_remove = []
                                        for key in equipment['photos'].keys():
                                            if 'marvel_' in key:
                                                keys_to_remove.append(key)
                                        for key in keys_to_remove:
                                            del equipment['photos'][key]
                                    
                                    # Clear Marvel-related session state keys
                                    keys_to_clear = []
                                    for key in st.session_state.keys():
                                        if f"marvel_" in key and equipment_key_prefix in key:
                                            keys_to_clear.append(key)
                                    for key in keys_to_clear:
                                        del st.session_state[key]
                            
                            with col2:
                                location_key = f"location_{equipment_key_prefix}"
                                # Initialize if not exists
                                if location_key not in st.session_state:
                                    st.session_state[location_key] = equipment.get('location', '')
                                    
                                st.text_input(
                                    "Location*",
                                    key=location_key,
                                    placeholder="e.g., Near entrance, Back wall, etc."
                                )
                                # Update the equipment data from session state
                                equipment['location'] = st.session_state[location_key]
                            
                            # If equipment type is selected, show checklist
                            if equipment['type']:
                                st.markdown("##### Inspection Checklist")
                                equipment_config = EQUIPMENT_TYPES[equipment['type']]
                                
                                # Render checklist items with full conditional logic
                                for i, item in enumerate(equipment_config['checklist']):
                                    if i > 0:
                                        st.markdown("<hr style='margin: 0.5rem 0;'>", unsafe_allow_html=True)  # Thinner separator
                                    render_checklist_item(equipment, item, equipment_key_prefix)
                                
                                # If "With Marvel" is checked, add Marvel checklist
                                if equipment.get('with_marvel', False):
                                    st.markdown("##### Marvel System Checklist")
                                    for i, item in enumerate(MARVEL_CHECKLIST):
                                        if i > 0:
                                            st.markdown("<hr style='margin: 0.5rem 0;'>", unsafe_allow_html=True)
                                        # Add prefix to distinguish Marvel questions
                                        render_checklist_item(equipment, item, equipment_key_prefix, prefix="marvel_")
                            
                            # Add separator between equipment
                            if equip_idx < len(kitchen['equipment_list']) - 1:
                                st.markdown("---")
        
    # Work Performed Section - Different for each report type
    if report_type == "General Service Report":
        st.markdown("### Work Performed")
        
        # Initialize work performed list in session state if not exists
        if 'work_performed_list' not in st.session_state:
            st.session_state.work_performed_list = []
        
        # Initialize a counter for work items to ensure unique keys
        if 'work_performed_counter' not in st.session_state:
            st.session_state.work_performed_counter = 0
        
        # Function to add a work performed item
        def add_work_performed():
            import time
            unique_id = f"{st.session_state.work_performed_counter}_{int(time.time() * 1000)}"
            work_number = len(st.session_state.work_performed_list) + 1
            st.session_state.work_performed_list.append({
                'id': unique_id,
                'title': f'Work Performed {work_number}',
                'description': '',
                'photos': [],
                'photo_descriptions': {}
            })
            st.session_state.work_performed_counter += 1
        
        # Function to remove a work performed item
        def remove_work_performed(index):
            st.session_state.work_performed_list.pop(index)
        
        # Add new work performed button
        col_add1, col_add2 = st.columns([1, 3])
        with col_add1:
            if st.button("➕ Add Work Item", type="secondary", key="add_work_performed_btn"):
                add_work_performed()
                st.rerun()
        
        # Display existing work performed items
        items_to_remove = []
        for i, work_item in enumerate(st.session_state.work_performed_list):
            with st.expander(f"Work Item {i+1}", expanded=True):
                # Work title
                title_key = f"work_title_{work_item['id']}"
                if title_key not in st.session_state:
                    st.session_state[title_key] = work_item.get('title', f'Work Performed {i+1}')
                
                title = st.text_input(
                    f"Work Title {i+1}",
                    key=title_key,
                    placeholder="Title of work performed"
                )
                work_item['title'] = st.session_state[title_key]
                
                # Work description
                desc_key = f"work_desc_{work_item['id']}"
                if desc_key not in st.session_state:
                    st.session_state[desc_key] = work_item.get('description', '')
                
                description = st.text_area(
                    f"Work Description {i+1}*",
                    key=desc_key,
                    placeholder="Describe the work performed (e.g., Replaced filters, Cleaned hood system, etc.)",
                    height=80
                )
                work_item['description'] = st.session_state[desc_key]
                
                # Photo upload for this work item
                st.markdown("#### Photos for this work item")
                uploaded_files = st.file_uploader(
                    f"Upload photos for Work Item {i+1}",
                    type=['png', 'jpg', 'jpeg'],
                    key=f"work_photos_{work_item['id']}",
                    accept_multiple_files=True
                )
                
                if uploaded_files:
                    work_item['photos'] = uploaded_files
                    
                    # Photo descriptions
                    st.markdown("##### Photo Descriptions")
                    for j, photo in enumerate(uploaded_files):
                        photo_desc_key = f"work_photo_desc_{work_item['id']}_{j}"
                        if photo_desc_key not in st.session_state:
                            # Auto-fill with default description if empty
                            default_desc = f"{work_item.get('title', f'Work Item {i+1}')} - Photo {j+1}"
                            st.session_state[photo_desc_key] = work_item.get('photo_descriptions', {}).get(str(j), default_desc)
                        
                        photo_desc = st.text_input(
                            f"Description for Photo {j+1}",
                            key=photo_desc_key,
                            placeholder="Brief description of what this photo shows"
                        )
                        work_item['photo_descriptions'][str(j)] = st.session_state[photo_desc_key]
                
                # Remove button
                col1, col2 = st.columns([3, 1])
                with col2:
                    if st.button(f"🗑️ Remove", key=f"remove_work_{work_item['id']}", type="secondary"):
                        items_to_remove.append(i)
        
        # Remove items marked for deletion
        if items_to_remove:
            for index in reversed(items_to_remove):
                remove_work_performed(index)
            st.rerun()
        
        # If no work items added, show a note
        if len(st.session_state.work_performed_list) == 0:
            st.info("No work items added. Click 'Add Work Item' to add work performed.")
    
    # Spare Parts Section (outside form to allow button interactions)
    st.markdown("### Spare Parts Required")
    
    # Initialize spare parts list in session state if not exists
    if 'spare_parts' not in st.session_state:
        st.session_state.spare_parts = []
    
    # Initialize a counter for spare parts to ensure unique keys
    if 'spare_parts_counter' not in st.session_state:
        st.session_state.spare_parts_counter = 0
    
    # Function to add a spare part row
    def add_spare_part():
        # Use timestamp to ensure truly unique IDs
        import time
        unique_id = f"{st.session_state.spare_parts_counter}_{int(time.time() * 1000)}"
        st.session_state.spare_parts.append({
            'id': unique_id,
            'name': '', 
            'quantity': 1
        })
        st.session_state.spare_parts_counter += 1
    
    # Function to remove a spare part row
    def remove_spare_part(index):
        st.session_state.spare_parts.pop(index)
    
    # Add new spare part button
    col_add1, col_add2 = st.columns([1, 3])
    with col_add1:
        if st.button("➕ Add Spare Part", type="secondary", key="add_spare_part_btn"):
            add_spare_part()
            st.rerun()
    
    # Helper functions for updating spare parts
    def update_part_name(index, part_id):
        if f"spare_part_name_{part_id}" in st.session_state:
            st.session_state.spare_parts[index]['name'] = st.session_state[f"spare_part_name_{part_id}"]
    
    def update_part_quantity(index, part_id):
        if f"spare_part_qty_{part_id}" in st.session_state:
            st.session_state.spare_parts[index]['quantity'] = st.session_state[f"spare_part_qty_{part_id}"]
    
    # Display existing spare parts
    parts_to_remove = []
    for i, part in enumerate(st.session_state.spare_parts):
        col1, col2, col3 = st.columns([3, 1, 0.5])
        with col1:
            # Use the part's ID for the key to ensure uniqueness
            # If part doesn't have an ID, create one
            if 'id' not in part:
                import time
                part['id'] = f"{st.session_state.spare_parts_counter}_{int(time.time() * 1000)}_{i}"
                st.session_state.spare_parts_counter += 1
            part_id = part['id']
            st.text_input(
                f"Spare Part {i+1} Name",
                value=part.get('name', ''),
                key=f"spare_part_name_{part_id}",
                placeholder="e.g., KSA Filter, UV Lamp, Solenoid Valve",
                on_change=update_part_name,
                args=(i, part_id)
            )
        with col2:
            st.number_input(
                f"Quantity",
                value=part.get('quantity', 1),
                min_value=1,
                key=f"spare_part_qty_{part_id}",
                on_change=update_part_quantity,
                args=(i, part_id)
            )
        with col3:
            st.markdown("<br>", unsafe_allow_html=True)  # Add spacing
            if st.button("🗑️", key=f"remove_part_{part_id}", help="Remove this spare part"):
                parts_to_remove.append(i)
    
    # Remove parts marked for deletion
    if parts_to_remove:
        for index in reversed(parts_to_remove):
            remove_spare_part(index)
        st.rerun()
    
    # If no spare parts added, show a note
    if len(st.session_state.spare_parts) == 0:
        st.info("No spare parts required. Click 'Add Spare Part' if parts are needed.")
    
    # Continue with the rest of the form
    with st.form("technical_report_form"):
        # Work Performed Section for Technical Report
        if report_type == "Technical Report":
            st.markdown("### Work Performed")
            work_performed = st.text_area(
                "Describe Work Performed*",
                placeholder="Detail all maintenance, repairs, or services completed...",
                height=120
            )
        
        # Recommendations Section
        st.markdown("### Recommendations")
        recommendations = st.text_area(
            "Recommendations",
            placeholder="Suggest any follow-up actions, parts needed, or future maintenance...",
            height=80
        )
        
        # Technician Information Section
        st.markdown("### Technician Information")
        col3, col4 = st.columns(2)
        
        with col3:
            technician_name = st.text_input("Technician Name*", placeholder="Enter your full name")
        
        with col4:
            service_date = st.date_input("Service Date", value=datetime.now())
        
        # Signature Section
        st.markdown("### Technician Signature")
        st.markdown("Please draw your signature below using your mouse or touchscreen")
        
        # Create signature canvas
        col1, col2 = st.columns([4, 1])
        
        with col1:
            canvas_result = st_canvas(
                fill_color="rgba(255, 255, 255, 0)",  # Transparent fill
                stroke_width=3,
                stroke_color="#000000",
                background_color="#FFFFFF",
                background_image=None,
                update_streamlit=True,
                height=120,
                width=450,
                drawing_mode="freedraw",
                point_display_radius=0,
                display_toolbar=True,
                key="signature_canvas",
            )
        
        with col2:
            st.markdown("### ")  # Add spacing
            st.info("Use the trash icon in the canvas toolbar to clear")
        
        # Check if signature is drawn
        if canvas_result is not None and canvas_result.image_data is not None:
            # Check if canvas has any drawing (non-transparent pixels)
            if np.any(canvas_result.image_data[:,:,3] > 0):
                st.success("✅ Technician signature captured")
        
        # Customer Signature Section
        st.markdown("### Customer Signature")
        col3, col4 = st.columns(2)
        
        with col3:
            # Customer name field - default to customer name from general info
            st.text_input(
                "Customer Representative Name",
                value=st.session_state.get('customer_name', ''),
                placeholder="Enter customer representative name",
                key="customer_signatory"
            )
        
        with col4:
            st.markdown("#### ")  # Spacing
            
        st.markdown("Please have the customer draw their signature below")
        
        # Create customer signature canvas
        col5, col6 = st.columns([4, 1])
        
        with col5:
            customer_canvas_result = st_canvas(
                fill_color="rgba(255, 255, 255, 0)",  # Transparent fill
                stroke_width=3,
                stroke_color="#000000",
                background_color="#FFFFFF",
                background_image=None,
                update_streamlit=True,
                height=120,
                width=450,
                drawing_mode="freedraw",
                point_display_radius=0,
                display_toolbar=True,
                key="customer_signature_canvas",
            )
        
        with col6:
            st.markdown("### ")  # Add spacing
            st.info("Use the trash icon in the canvas toolbar to clear")
        
        # Check if customer signature is drawn
        if customer_canvas_result is not None and customer_canvas_result.image_data is not None:
            # Check if canvas has any drawing (non-transparent pixels)
            if np.any(customer_canvas_result.image_data[:,:,3] > 0):
                st.success("✅ Customer signature captured")
        
        # Submit button
        submitted = st.form_submit_button("Generate Report", type="primary")
        
        if submitted:
            # Get values from session state
            customer_name = st.session_state.get('customer_name', '')
            project_name = st.session_state.get('project_name', '')
            contact_person = st.session_state.get('contact_person', '')
            outlet_location = st.session_state.get('outlet_location', '')
            contact_number = st.session_state.get('contact_number', '')
            visit_type = st.session_state.get('visit_type', '')
            visit_class = st.session_state.get('visit_class', '')
            date = st.session_state.get('report_date', datetime.now())
            
            # Validation - different based on report type
            if report_type == "Testing and Commissioning Report":
                required_fields = {
                    "Customer's Name": customer_name,
                    "Project Name": project_name,
                    "Contact Person": contact_person,
                    "Outlet/Location": outlet_location,
                    "Contact Number": contact_number,
                    "Visit Type": visit_type,
                    "Technician Name": technician_name
                }
            elif report_type == "Technical Report":
                required_fields = {
                    "Customer's Name": customer_name,
                    "Project Name": project_name,
                    "Contact Person": contact_person,
                    "Outlet/Location": outlet_location,
                    "Contact Number": contact_number,
                    "Visit Type": visit_type,
                    "Work Performed": work_performed,
                    "Spare Parts": st.session_state.get('spare_parts', []),
                    "Technician Name": technician_name
                }
            else:  # General Service Report
                # Check work performed list
                work_items_valid = True
                if not st.session_state.get('work_performed_list'):
                    work_items_valid = False
                else:
                    for item in st.session_state.work_performed_list:
                        if not item.get('description', '').strip():
                            work_items_valid = False
                            break
                
                required_fields = {
                    "Customer's Name": customer_name,
                    "Project Name": project_name,
                    "Contact Person": contact_person,
                    "Outlet/Location": outlet_location,
                    "Contact Number": contact_number,
                    "Visit Type": visit_type,
                    "Work Performed Items": "Valid" if work_items_valid else "",
                    "Spare Parts": st.session_state.get('spare_parts', []),
                    "Technician Name": technician_name
                }
            
            missing_fields = [field for field, value in required_fields.items() if not value]
            
            # Validate kitchen and equipment data (only for Technical Report)
            equipment_errors = []
            if report_type == "Technical Report":
                for kitchen_idx, kitchen in enumerate(st.session_state.kitchen_list):
                    kitchen_name = kitchen.get('name', f'Kitchen {kitchen_idx + 1}')
                    for equip_idx, equipment in enumerate(kitchen.get('equipment_list', [])):
                        equip_name = f"{kitchen_name} - Equipment #{equip_idx + 1}"
                        if not equipment.get('type'):
                            equipment_errors.append(f"{equip_name}: Equipment type is required")
                        elif not equipment.get('location'):
                            equipment_errors.append(f"{equip_name}: Location is required")
            
            # Show reminders for missing fields but don't block submission
            if missing_fields or equipment_errors:
                st.warning("📋 **Reminder:** The following fields are recommended but not required:")
                if missing_fields:
                    st.write("**General Information:**")
                    for field in missing_fields:
                        st.write(f"• {field}")
                if equipment_errors:
                    st.write("**Equipment Information:**")
                    for error in equipment_errors:
                        st.write(f"• {error}")
                st.info("The report will be generated with the available information.")
            
            # Always proceed with report generation
            # Process signature from canvas
            signature_img = None
            if canvas_result.image_data is not None and np.any(canvas_result.image_data[:,:,3] > 0):
                # Convert canvas to image
                sig_image = Image.fromarray(canvas_result.image_data.astype('uint8'), 'RGBA')
                
                # Create white background
                white_bg = Image.new('RGB', sig_image.size, 'white')
                white_bg.paste(sig_image, mask=sig_image.split()[3])
                
                # Find the bounding box of the signature
                bbox = white_bg.getbbox()
                if bbox:
                    # Crop to signature
                    cropped = white_bg.crop(bbox)
                    
                    # Resize if too large
                    max_width, max_height = 200, 80
                    cropped.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                    
                    # Convert to bytes
                    img_bytes = io.BytesIO()
                    cropped.save(img_bytes, format='PNG')
                    img_bytes.seek(0)
                    signature_img = img_bytes
            
            # Process customer signature from canvas
            customer_signature_img = None
            if customer_canvas_result.image_data is not None and np.any(customer_canvas_result.image_data[:,:,3] > 0):
                # Convert canvas to image
                cust_sig_image = Image.fromarray(customer_canvas_result.image_data.astype('uint8'), 'RGBA')
                
                # Create white background
                white_bg = Image.new('RGB', cust_sig_image.size, 'white')
                white_bg.paste(cust_sig_image, mask=cust_sig_image.split()[3])
                
                # Find the bounding box of the signature
                bbox = white_bg.getbbox()
                if bbox:
                    # Crop to signature
                    cropped = white_bg.crop(bbox)
                    
                    # Resize if too large
                    max_width, max_height = 200, 80
                    cropped.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                    
                    # Convert to bytes
                    img_bytes = io.BytesIO()
                    cropped.save(img_bytes, format='PNG')
                    img_bytes.seek(0)
                    customer_signature_img = img_bytes
            
            # Collect all data based on report type
            if report_type == "Testing and Commissioning Report":
                report_data = {
                    'report_type': report_type,
                    'customer_name': customer_name,
                    'project_name': project_name,
                    'contact_person': contact_person,
                    'outlet_location': outlet_location,
                    'contact_number': contact_number,
                    'visit_type': visit_type,
                    'visit_class': visit_class,
                    'date': date.strftime('%Y-%m-%d'),
                    'canopy_data': st.session_state.get('canopy_data', []),
                    'tc_checklists': st.session_state.get('tc_checklists', {}),
                    'recommendations': recommendations,
                    'technician_name': technician_name,
                    'service_date': service_date.strftime('%Y-%m-%d'),
                    'technician_signature': signature_img
                }
            elif report_type == "Technical Report":
                report_data = {
                    'report_type': report_type,
                    'customer_name': customer_name,
                    'project_name': project_name,
                    'contact_person': contact_person,
                    'outlet_location': outlet_location,
                    'contact_number': contact_number,
                    'visit_type': visit_type,
                    'visit_class': visit_class,
                    'date': date.strftime('%Y-%m-%d'),
                    'equipment_inspection': get_kitchen_summary(),
                    'kitchen_list': st.session_state.kitchen_list,
                    'work_performed': work_performed,
                    'spare_parts': st.session_state.get('spare_parts', []),
                    'recommendations': recommendations,
                    'technician_name': technician_name,
                    'service_date': service_date.strftime('%Y-%m-%d'),
                    'technician_signature': signature_img,
                    'customer_signatory': st.session_state.get('customer_signatory', customer_name),
                    'customer_signature': customer_signature_img
                }
            else:  # General Service Report
                report_data = {
                    'report_type': report_type,
                    'customer_name': customer_name,
                    'project_name': project_name,
                    'contact_person': contact_person,
                    'outlet_location': outlet_location,
                    'contact_number': contact_number,
                    'visit_type': visit_type,
                    'visit_class': visit_class,
                    'date': date.strftime('%Y-%m-%d'),
                    'work_performed_list': st.session_state.get('work_performed_list', []),
                    'spare_parts': st.session_state.get('spare_parts', []),
                    'recommendations': recommendations,
                    'technician_name': technician_name,
                    'service_date': service_date.strftime('%Y-%m-%d'),
                    'technician_signature': signature_img,
                    'customer_signatory': st.session_state.get('customer_signatory', customer_name),
                    'customer_signature': customer_signature_img
                }
            
            # Store data in session state for download outside form
            st.session_state.report_data = report_data
            st.session_state.report_generated = True
            st.session_state.saved_customer_name = customer_name
            st.session_state.saved_report_date = date
    
    # Form sharing section (outside of form)
    st.markdown("### 🔗 Share Form Data")
    st.markdown("Save your current form inputs as a shareable link. Photos and signatures are not included.")
    
    col_share1, col_share2 = st.columns([1, 1])
    with col_share1:
        if st.button("🔗 Generate Shareable Link", type="secondary"):
            share_url = generate_shareable_link()
            if share_url:
                st.session_state['generated_share_url'] = share_url
                st.success("Link generated successfully!")
            else:
                st.error("Failed to generate shareable link")
    
    with col_share2:
        if st.session_state.get('generated_share_url'):
            st.text_area(
                "Shareable Link (Copy and share):",
                value=st.session_state['generated_share_url'],
                height=100,
                help="Copy this URL and share it. When opened, it will prefill the form with current data."
            )
    
    # Handle report download outside of form
    if st.session_state.get('report_generated', False):
        try:
            # Generate the report based on type
            report_type = st.session_state.report_data.get('report_type', 'Technical Report')
            
            if report_type == "Technical Report":
                doc_bytes = create_technical_report(st.session_state.report_data)
                filename_prefix = "Technical_Report"
            elif report_type == "Testing and Commissioning Report":
                doc_bytes = create_testing_commissioning_report(st.session_state.report_data)
                filename_prefix = "Testing_Commissioning_Report"
            else:  # General Service Report
                doc_bytes = create_general_service_report(st.session_state.report_data)
                filename_prefix = "General_Service_Report"
            
            # Create filename
            customer_name = st.session_state.saved_customer_name
            date = st.session_state.saved_report_date
            filename = f"{filename_prefix}_{customer_name.replace(' ', '_')}_{date.strftime('%Y%m%d')}.docx"
            
            # Success message
            st.success("✅ Report generated successfully!")
            
            # Download button (outside form)
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                st.download_button(
                    label="📥 Download Report",
                    data=doc_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            # Option to generate another report
            if st.button("Generate Another Report", type="secondary"):
                # Clear all session state data for a fresh start
                st.session_state.report_generated = False
                st.session_state.kitchen_list = []
                st.session_state.report_data = {}
                # Clear all widget keys
                keys_to_clear = []
                for key in st.session_state.keys():
                    if (key.startswith('q_') or key.startswith('comment_') or 
                        key.startswith('equip_type_') or key.startswith('with_marvel_') or 
                        key.startswith('location_') or key.startswith('photo_') or
                        key == 'customer_signatory' or key == 'customer_signature_canvas' or
                        key == 'signature_canvas'):
                        keys_to_clear.append(key)
                for key in keys_to_clear:
                    del st.session_state[key]
                st.rerun()
                
        except Exception as e:
            st.error(f"❌ Error generating report: {str(e)}")
            if st.button("Try Again"):
                st.session_state.report_generated = False
                st.rerun()

if __name__ == "__main__":
    main()