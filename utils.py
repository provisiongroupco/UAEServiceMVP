"""
Utility functions for the Service Reports application
"""

from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# Professional brand colors
HALTON_BLUE = RGBColor(31, 71, 136)  # #1f4788
HALTON_LIGHT_BLUE = RGBColor(44, 90, 160)  # #2c5aa0
HALTON_DARK_GRAY = RGBColor(64, 64, 64)  # #404040

def add_header_with_logo(doc, logo_path=None):
    """
    Add a professional header with company branding
    
    Args:
        doc: Document object
        logo_path: Path to logo file (optional)
    """
    header = doc.sections[0].header
    
    # Create a table for header layout
    header_table = header.add_table(rows=1, cols=3, width=Inches(6.5))
    header_table.autofit = False
    
    # Left cell for logo (placeholder text for now)
    left_cell = header_table.cell(0, 0)
    left_cell.width = Inches(2)
    if logo_path:
        # Add logo when available
        left_cell.paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.5))
    else:
        # Placeholder text
        logo_text = left_cell.paragraphs[0].add_run("COMPANY")
        logo_text.font.size = Pt(20)
        logo_text.font.bold = True
        logo_text.font.color.rgb = HALTON_BLUE
    
    # Middle cell for company name
    middle_cell = header_table.cell(0, 1)
    middle_cell.width = Inches(3.5)
    company_para = middle_cell.paragraphs[0]
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_text = company_para.add_run("SERVICE DIVISION\nTechnical Services")
    company_text.font.size = Pt(12)
    company_text.font.color.rgb = HALTON_DARK_GRAY
    
    # Right cell for contact info
    right_cell = header_table.cell(0, 2)
    right_cell.width = Inches(1)
    contact_para = right_cell.paragraphs[0]
    contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact_text = contact_para.add_run("Service Department")
    contact_text.font.size = Pt(10)
    contact_text.font.color.rgb = HALTON_DARK_GRAY
    
    # Add a line below header
    add_horizontal_line(header)

def add_horizontal_line(parent):
    """Add a horizontal line to separate sections"""
    p = parent.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pBorder = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1f4788')
    pBorder.append(bottom)
    p._p.get_or_add_pPr().append(pBorder)

def format_table_style(table):
    """Apply professional branding to tables"""
    # Set table borders
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)
    
    tbl.tblPr.append(tblBorders)
    
    # Format header row if exists
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = HALTON_BLUE

def add_footer(doc):
    """Add a professional footer"""
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = footer_para.add_run("Technical Service Report - Confidential")
    footer_text.font.size = Pt(8)
    footer_text.font.color.rgb = HALTON_DARK_GRAY
    
    # Add page numbers
    footer_para.add_run(" | Page ")
    add_page_number(footer_para)

def add_page_number(paragraph):
    """Add page number field to paragraph"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_cell_margins(cell, top=0, bottom=0, left=0.1, right=0.1):
    """Set cell margins"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for margin_type, value in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        margin = OxmlElement(f'w:{margin_type}')
        margin.set(qn('w:w'), str(int(value * 1440)))  # Convert inches to twips
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    
    tcPr.append(tcMar)

def style_heading(heading, level=1):
    """Apply consistent styling to headings"""
    if level == 0:  # Title
        heading.style.font.size = Pt(20)
        heading.style.font.color.rgb = HALTON_BLUE
        heading.style.font.bold = True
    elif level == 1:  # Main sections
        heading.style.font.size = Pt(14)
        heading.style.font.color.rgb = HALTON_BLUE
        heading.style.font.bold = True
    else:  # Subsections
        heading.style.font.size = Pt(12)
        heading.style.font.color.rgb = HALTON_DARK_GRAY
        heading.style.font.bold = True

def create_info_table(doc, data_rows, col_widths=[2.5, 4]):
    """Create a professionally formatted information table"""
    table = doc.add_table(rows=len(data_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for i, width in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)
    
    # Fill table with data
    for i, (label, value) in enumerate(data_rows):
        # Label cell
        label_cell = table.cell(i, 0)
        label_cell.text = label
        label_para = label_cell.paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_run = label_para.runs[0]
        label_run.font.bold = True
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = HALTON_DARK_GRAY
        
        # Value cell
        value_cell = table.cell(i, 1)
        value_cell.text = str(value) if value else ""
        value_para = value_cell.paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        value_run = value_para.runs[0]
        value_run.font.size = Pt(11)
        
        # Set cell margins
        set_cell_margins(label_cell)
        set_cell_margins(value_cell)
        
        # Add subtle shading to alternate rows
        if i % 2 == 1:
            for cell in [label_cell, value_cell]:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F5F5F5')
                tcPr.append(shd)
    
    # Apply table borders
    format_table_style_enhanced(table)
    
    return table

def format_table_style_enhanced(table):
    """Apply enhanced professional branding to tables"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Remove existing borders
    for child in tblPr:
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
    
    # Create new borders
    tblBorders = OxmlElement('w:tblBorders')
    
    # Define border styles
    border_styles = {
        'top': {'sz': '12', 'color': '1f4788'},
        'bottom': {'sz': '12', 'color': '1f4788'},
        'left': {'sz': '4', 'color': 'CCCCCC'},
        'right': {'sz': '4', 'color': 'CCCCCC'},
        'insideH': {'sz': '4', 'color': 'E0E0E0'},
        'insideV': {'sz': '0', 'color': 'FFFFFF'}  # No vertical lines
    }
    
    for border_name, style in border_styles.items():
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), style['sz'])
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), style['color'])
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def add_logo_to_doc(doc, logo_path=None):
    """Check for logo in assets folder and add to document"""
    if not logo_path:
        # Check for logo in assets folder
        possible_logos = ['halton_logo.png', 'halton_logo.jpg', 'logo.png', 'logo.jpg']
        for logo_name in possible_logos:
            check_path = os.path.join('assets', logo_name)
            if os.path.exists(check_path):
                logo_path = check_path
                break
    
    return logo_path